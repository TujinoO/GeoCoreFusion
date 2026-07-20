"""Build the V6.1 visual-full-detail DOCX from the retained V5 template.

The retained reference is copied before its main body is replaced.  Styles,
theme, numbering foundation, header/footer parts, section geometry, custom XML,
and source media remain the design authority.  This builder performs structural
audits only; the resulting DOCX must still pass Word/PDF/PNG page-by-page QA.

No file or directory is deleted by this script.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import shutil
import zipfile
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as _Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Pt

try:
    import build_v6_document as base
except ModuleNotFoundError:
    from scripts import build_v6_document as base


REFERENCE_SHA256 = base.REFERENCE_SHA256
DEFAULT_REFERENCE = Path(
    r"E:\GRP Docs\2026.07 岩心近景影像融合技术"
    r"\高水平论文研究路线与实验方案_V5最终结果补充版.docx"
)
OUTPUT_NAME = "高水平论文研究路线与实验方案_V6.1视觉全细节融合版.docx"


FIGURE_SPECS: dict[str, dict[str, Any]] = {
    "01_v61_visual_full_detail_route.png": {
        "title": "V6.1 视觉全细节融合总路线：配准残差控制、RGB 细节重建与观测回投",
        "alt": (
            "从配准残差前置控制、RGB 双边降噪和四尺度分解，到亮度色度最强梯度、"
            "screened-Poisson 重建、全细节注入及最终产品低分辨率回投的技术路线图。"
        ),
    },
    "02_3dssz_v6_v61_shared_domain_comparison.png": {
        "title": "3DSSZ：V6 与 V6.1 的共享反射率显示域全图、纹理区和暗区对比",
        "alt": (
            "3DSSZ 的 RGB、V6、V6.1 和反射率差异图，包含一处岩心纹理区和一处暗区局部放大。"
        ),
        "max_height_in": 5.45,
    },
    "03_zkh3_v6_v61_shared_domain_comparison.png": {
        "title": "ZKH3：V6 与 V6.1 的共享反射率显示域全图、纹理区和暗区对比",
        "alt": (
            "ZKH3 的 RGB、V6、V6.1 和反射率差异图，包含碎裂岩心纹理区和暗区局部放大。"
        ),
    },
    "04_v6_v61_key_metric_improvement.png": {
        "title": "V6 到 V6.1 的关键指标：边界、可靠/暗区细节、edge F1 与最终产品回投",
        "alt": (
            "两场景 V6 和 V6.1 的 RGB 材料边界相关、2201 nm 可靠细节和暗区可靠纹理相关、"
            "最终产品回投 RMSE、SAM 与 edge F1 分组柱状图。"
        ),
    },
    "05_v61_detail_amplitude_and_residual_audit.png": {
        "title": "V6.1 全细节注入的幅值与非相干高频审计",
        "alt": (
            "两场景三个目标波段的 rho、beta、高频能量比 A 和正交残差 R-perp 的 V6/V6.1 对比图。"
        ),
        "note": (
            "这些指标使用同一 RGB 作为细节引导和结构参考；它们衡量空间迁移、幅值及伪影风险，"
            "不构成独立 HR-SWIR 真值。"
        ),
    },
    "06_v61_low_resolution_spectral_curve_consistency.png": {
        "title": "最终 HR 产品回投到原传感器网格后的代表性光谱曲线一致性",
        "alt": (
            "3DSSZ 与 ZKH3 暗、中、亮反射率代表性低分辨率像素的原始观测曲线和 V6.1 HR 产品 PSF 回投曲线。"
        ),
        "page_break_before": True,
    },
    "07_v61_registration_residual_control.png": {
        "title": "配准残差前置控制、真实 ROI 结构得分与合成真值亚像素证据",
        "alt": (
            "配准处理链、两场景 NIR/SWIR ROI 修正前后结构得分、以及受控合成 TRE/EPE 的 Median 与 P95。"
        ),
    },
    "08_v61_single_product_and_claim_contract.png": {
        "title": "V6.1 单产品合同与可验证/不可验证主张边界",
        "alt": (
            "原始 RGB 和低分辨率 NIR/SWIR 进入 V6.1 融合算子后生成唯一视觉全细节产品，"
            "并区分可验证的观测一致性与尚不可证明的独立 HR-SWIR 真值。"
        ),
    },
}


OVERVIEW_ITEMS = [
    "版本决策：回到 V6 直观全细节路线，V8 条件门控不进入当前产品",
    "单产品合同：全部降噪、共配准 RGB 空间细节允许进入 NIR/SWIR",
    "配准前端：坐标一致、无效边界、ROI 推扫几何与双向亚像素同名点",
    "细节提取：双边降噪、四尺度金字塔和亮度/色度最强梯度",
    "可重建细节：screened-Poisson 消除梯度不可积和多尺度相位冲突",
    "注入与锐化：log 增益、受限加性细节、暗区对比和 halo-limited sharpen",
    "光谱闭环：实际最终 HR 产品经 PSF 回投到原 NIR/SWIR 观测网格",
    "真实 ROI：3DSSZ/ZKH3 共享反射率显示域、局部放大和差异图",
    "量化审计：rho、beta、A、R-perp、dark/flat、edge F1、RMSE/SAM",
    "论文路线：独立 TRE、MTF/暗场标定、跨钻孔盲测和局部 HR 光谱真值",
]


CODE_MAP_ROWS = [
    ["模块", "主要文件", "V6.1 职责", "验证入口"],
    [
        "配准前端",
        "src/geocorefusion/registration.py",
        "坐标端点、有效边界、ROI 推扫几何、双向亚像素同名点",
        "tests/test_registration.py；图 7；registration_model.json",
    ],
    [
        "全细节提取",
        "src/geocorefusion/fusion.py",
        "RGB 去噪、四尺度细节、最强梯度和 screened-Poisson",
        "tests/test_fusion.py；图 1/2/3",
    ],
    [
        "注入与回投",
        "src/geocorefusion/fusion.py",
        "log 增益、加性细节、暗区对比、最终产品观测回投",
        "quality_report.json；图 4/6",
    ],
    [
        "产品合同",
        "src/geocorefusion/output.py",
        "geocorefusion.visual-full-detail.v1 单产品和非真值边界",
        "tests/test_output.py；manifest.json；图 8",
    ],
    [
        "质量评价",
        "src/geocorefusion/quality.py",
        "rho/beta/A/R-perp、dark/flat、edge F1、halo 和最终回投",
        "tests/test_quality_detail.py；图 4/5/6",
    ],
    [
        "科研图复现",
        "scripts/build_v61_figures.py",
        "同一反射率显示域、局部放大、光谱曲线和机器摘要",
        "figure_manifest.json；benchmark_summary.json",
    ],
]


PARAMETER_ROWS = [
    ["参数组", "冻结值", "作用"],
    ["RGB 去噪", "0.55", "抑制随机亮度/色度噪声，保留边缘"],
    ["金字塔 σ / px", "0.65 / 1.35 / 2.80 / 5.60", "覆盖颗粒、裂隙、层理和暗区缓变对比"],
    ["金字塔权重", "1.15 / 1.00 / 0.82 / 0.58", "补偿细尺度衰减，限制大尺度压制"],
    ["色度权重", "0.35", "恢复等亮彩色边缘"],
    ["Poisson screen", "0.24", "保持四尺度先验并避免低频漂移"],
    ["暗区增强", "0.25", "增强暗区内部相对对比"],
    ["log 注入", "0.92；gain 0.52–1.92", "强空间细节和有限幅值"],
    ["加性细节", "0.22", "补充低反射率和色彩边缘绝对幅值"],
    ["最终回投", "4 次；weight 0.80；clip 0.65σ", "控制实际最终产品的 LR 观测残差"],
    ["锐化", "0.25", "有限边缘恢复，不继续盲目增大"],
]


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def replace_header_version(document: _Document) -> None:
    replacement = (
        "GeoCoreFusion V6.1    岩心 RGB-NIR-SWIR 配准与视觉全细节融合研究方案"
    )
    for section in document.sections:
        header = section.header
        changed = False
        for paragraph in header.paragraphs:
            text_nodes = list(paragraph._p.iter(qn("w:t")))
            visible = "".join(node.text or "" for node in text_nodes)
            if re.search(r"V(?:5|6(?:\.\d+)?|7|8)", visible) and text_nodes:
                text_nodes[0].text = replacement
                for node in text_nodes[1:]:
                    node.text = ""
                changed = True
        if changed:
            continue
        paragraph = next((p for p in header.paragraphs if p.text.strip()), None)
        if paragraph is None:
            raise RuntimeError("retained running header has no editable paragraph")
        if paragraph.runs:
            paragraph.runs[0].text = replacement
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(replacement)


def add_cover(document: _Document, build_date: str) -> None:
    spacer = document.add_paragraph(style="Normal")
    spacer.paragraph_format.space_after = Pt(54)

    eyebrow = document.add_paragraph(style="Normal")
    eyebrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eyebrow.paragraph_format.space_after = Pt(16)
    run = eyebrow.add_run("RESEARCH ROADMAP & EXPERIMENTAL PLAN")
    base.set_run_font(run, latin="Calibri", size=10.5, color=base.BLUE, bold=True)

    title = document.add_paragraph(style="Normal")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    title.paragraph_format.line_spacing = 1.04
    run = title.add_run("岩心 RGB–NIR–SWIR 近景影像配准与融合")
    base.set_run_font(run, latin="Calibri", size=28, color=base.NAVY, bold=True)
    run = title.add_run("\nV6.1 视觉全细节融合研究方案")
    base.set_run_font(run, latin="Calibri", size=24, color=base.NAVY, bold=True)

    subtitle = document.add_paragraph(style="Normal")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    run = subtitle.add_run(
        "亚像素能力配准 · 去噪 RGB 全细节 · screened-Poisson · 最终产品观测回投"
    )
    base.set_run_font(run, latin="Calibri", size=14, color=base.MID_BLUE)

    project = document.add_paragraph(style="Normal")
    project.alignment = WD_ALIGN_PARAGRAPH.CENTER
    project.paragraph_format.space_after = Pt(7)
    run = project.add_run("GeoCoreFusion | RGB–NIR–SWIR Drill-Core Image Fusion")
    base.set_run_font(run, latin="Calibri", size=11, color=base.GRAY, bold=True)

    date_paragraph = document.add_paragraph(style="Normal")
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_paragraph.add_run(f"{build_date} | V6.1 视觉全细节融合版")
    base.set_run_font(run, latin="Calibri", size=10.5, color=base.GRAY)
    date_paragraph.paragraph_format.space_after = Pt(36)

    status = document.add_paragraph(style="Normal")
    status.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = status.add_run(
        "冻结结论：两场景空间细节、暗区可靠纹理与最终产品光谱回投均优于 V6"
    )
    base.set_run_font(run, latin="Calibri", size=9.5, color=base.ORANGE, bold=True)
    status.paragraph_format.space_after = Pt(8)

    document.add_paragraph(style="Normal").add_run().add_break(WD_BREAK.PAGE)


def add_overview(document: _Document) -> None:
    heading = document.add_paragraph("内容概览", style="Heading 1")
    base.set_keep(heading, next_=True)
    lead = document.add_paragraph(
        "V6.1 只保留一条视觉全细节产品路线：在最新配准前端和 RGB 保边降噪之后，"
        "接受全部 RGB 空间细节；同时通过最终产品 PSF 回投保持原始尺度 NIR/SWIR 光谱观测一致性。",
        style="Normal",
    )
    lead.paragraph_format.space_after = Pt(7)

    numbering_id = base.new_numbering_instance(document)
    for text in OVERVIEW_ITEMS:
        paragraph = document.add_paragraph(style="List Number")
        base.apply_numbering(paragraph, numbering_id)
        paragraph.paragraph_format.space_after = Pt(3)
        base.add_rich_text(
            paragraph,
            text,
            base_size=10.5,
            base_color=base.MID_BLUE,
        )
        for run in paragraph.runs:
            run.bold = True

    toc_heading = document.add_paragraph(style="Normal")
    run = toc_heading.add_run("自动目录")
    base.set_run_font(run, latin="Calibri", size=13, color=base.BLUE, bold=True)
    toc_heading.paragraph_format.space_before = Pt(9)
    toc_heading.paragraph_format.space_after = Pt(4)
    base.set_keep(toc_heading, next_=True)

    toc = document.add_paragraph(style="Normal")
    base.add_complex_field(
        toc,
        r'TOC \o "1-2" \h \z \u',
        "目录域将在 Microsoft Word 打开或打印时更新。",
        font_size=9,
        color=base.GRAY,
    )
    note = document.add_paragraph(
        "提示：最终页码以 Word 字段更新和逐页视觉质检结果为准。",
        style="Figure Caption",
    )
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    document.add_paragraph(style="Normal").add_run().add_break(WD_BREAK.PAGE)


def metric_rows(benchmark: dict[str, Any]) -> list[list[str]]:
    rows = [
        [
            "场景",
            "方法",
            "边界相关",
            "2201 可靠 rho",
            "2201 暗区 rho",
            "edge F1",
            "回投 RMSE",
            "SAM / °",
            "band-CC",
        ]
    ]
    for scene in ("3dssz", "zkh3"):
        for method in ("V6", "V6.1"):
            metrics = benchmark["scenes"][scene]["methods"][method]["metrics"]
            band = metrics["bands"]["2201.0nm"]
            rows.append(
                [
                    scene.upper(),
                    method,
                    f"{metrics['boundary_correlation']:.3f}",
                    f"{band['reliable_rgb_detail']['rho']:.3f}",
                    f"{band['dark_reliable_rgb_detail']['rho']:.3f}",
                    f"{band['edge_f1_1px']:.3f}",
                    f"{metrics['forward_rmse']:.5f}",
                    f"{metrics['forward_sam_deg']:.3f}",
                    f"{metrics['forward_band_cc']:.5f}",
                ]
            )
    return rows


def add_traceability_appendices(state: base.BuildState) -> None:
    heading = state.document.add_paragraph(
        "附录 A：V6.1 代码—配置—证据地图", style="Heading 1"
    )
    heading.paragraph_format.page_break_before = True
    base.set_keep(heading, next_=True)
    base.add_table(state, CODE_MAP_ROWS, caption="V6.1 代码—配置—证据追溯地图")
    state.document.add_paragraph(style="Normal").add_run().add_break(WD_BREAK.PAGE)
    base.add_table(state, PARAMETER_ROWS, caption="V6.1 冻结参数及其作用")

    heading = state.document.add_paragraph(
        "附录 B：机器结果快照", style="Heading 1"
    )
    heading.paragraph_format.page_break_before = True
    base.set_keep(heading, next_=True)
    base.add_callout(
        state.document,
        "来源与口径",
        (
            "本附录直接读取 geocorefusion.v61-visual-benchmark.v1 JSON。"
            "V6/V6.1 显示图在同一场景反射率范围内映射；rho/edge 是同数据 RGB 引导诊断，"
            "RMSE/SAM/band-CC 是最终 HR 产品降采样后的原始尺度观测一致性。"
        ),
    )
    base.add_table(state, metric_rows(state.benchmark), caption="两组真实 ROI 的 V6/V6.1 机器结果")
    base.add_callout(
        state.document,
        "强制主张边界",
        (
            "共享显示域和 RGB/fused-band 指标不能证明独立 HR-SWIR 真值；"
            "受控合成 TRE/EPE 不能替代生产岩心的独立靶标 TRE。"
        ),
        warning=True,
    )


def validate_inputs(
    reference: Path,
    content: Path,
    benchmark: Path,
    figures_dir: Path,
    output: Path,
) -> None:
    for path in (reference, content, benchmark):
        if not path.is_file():
            raise FileNotFoundError(path)
    if not figures_dir.is_dir():
        raise NotADirectoryError(figures_dir)
    missing = [name for name in FIGURE_SPECS if not (figures_dir / name).is_file()]
    if missing:
        raise FileNotFoundError(f"missing V6.1 figures: {missing}")
    if reference.resolve() == output.resolve():
        raise ValueError("output must differ from retained V5 reference")
    actual = sha256_file(reference).lower()
    if actual != REFERENCE_SHA256:
        raise RuntimeError(
            "retained V5 SHA-256 mismatch; re-distill before authoring: "
            f"expected {REFERENCE_SHA256}, got {actual}"
        )
    payload = json.loads(benchmark.read_text(encoding="utf-8"))
    if payload.get("schema") != "geocorefusion.v61-visual-benchmark.v1":
        raise ValueError("benchmark is not the V6.1 shared-domain artifact")
    if payload.get("scientific_conditional_included") is not False:
        raise ValueError("V6.1 benchmark unexpectedly includes scientific_conditional")
    markdown = content.read_text(encoding="utf-8")
    referenced = {
        Path(match).name
        for match in re.findall(r"!\[[^\]]*\]\(([^)\s]+\.png)", markdown)
    }
    if unknown := sorted(referenced - set(FIGURE_SPECS)):
        raise ValueError(f"content references figures without metadata: {unknown}")


def package_foundation(reference: Path, output: Path) -> dict[str, int]:
    required_parts = {
        "word/styles.xml",
        "word/theme/theme1.xml",
        "word/fontTable.xml",
        "word/numbering.xml",
        "word/header1.xml",
        "word/footer1.xml",
        "word/settings.xml",
        "word/webSettings.xml",
    }
    with zipfile.ZipFile(reference) as source_zip, zipfile.ZipFile(output) as final_zip:
        source_names = set(source_zip.namelist())
        final_names = set(final_zip.namelist())
        if missing := sorted(required_parts - final_names):
            raise RuntimeError(f"V6.1 lost retained package parts: {missing}")
        source_custom = {name for name in source_names if name.startswith("customXml/")}
        if lost := sorted(source_custom - final_names):
            raise RuntimeError(f"V6.1 lost retained customXml parts: {lost}")
        source_media = {name for name in source_names if name.startswith("word/media/")}
        if lost := sorted(source_media - final_names):
            raise RuntimeError(f"V6.1 lost retained source media parts: {lost}")
    return {
        "required_foundation_parts": len(required_parts),
        "retained_custom_xml_parts": len(source_custom),
        "retained_source_media_parts": len(source_media),
    }


def expected_markdown_headings(markdown: str) -> list[str]:
    headings: list[tuple[int, str]] = []
    for line in markdown.splitlines():
        match = re.match(r"^(#{1,3})\s+(.+?)\s*$", line.strip())
        if match:
            headings.append((len(match.group(1)), match.group(2).strip()))
    for index, (level, _) in enumerate(headings):
        if level == 1:
            headings.pop(index)
            break
    return [text for _, text in headings]


def audit_document(path: Path, reference: Path, markdown: str) -> dict[str, Any]:
    document = Document(path)
    if len(document.sections) != 1:
        raise RuntimeError(f"expected one retained section, found {len(document.sections)}")
    section = document.sections[0]
    if section.orientation != WD_ORIENT.PORTRAIT:
        raise RuntimeError("V6.1 document is not portrait")
    page_width = section.page_width / 914400
    page_height = section.page_height / 914400
    if abs(page_width - 8.5) > 0.01 or abs(page_height - 11.0) > 0.01:
        raise RuntimeError(f"unexpected page size: {page_width:.3f} x {page_height:.3f} in")
    margins = [
        section.top_margin / 914400,
        section.right_margin / 914400,
        section.bottom_margin / 914400,
        section.left_margin / 914400,
    ]
    if any(abs(value - 1.0) > 0.01 for value in margins):
        raise RuntimeError(f"retained margins changed: {margins}")
    if not section.different_first_page_header_footer:
        raise RuntimeError("different-first-page setting was lost")

    headings = [
        p for p in document.paragraphs if p.style and p.style.name.startswith("Heading ")
    ]
    expected_counts = Counter(expected_markdown_headings(markdown))
    actual_counts = Counter(p.text.strip() for p in headings)
    if missing := list((expected_counts - actual_counts).elements()):
        raise RuntimeError(f"Markdown headings were not emitted as Word headings: {missing}")
    if len(headings) < 25:
        raise RuntimeError(f"unexpectedly few real headings: {len(headings)}")
    if len(document.inline_shapes) != len(FIGURE_SPECS):
        raise RuntimeError(
            f"expected {len(FIGURE_SPECS)} inline figures, found {len(document.inline_shapes)}"
        )
    for index, shape in enumerate(document.inline_shapes, start=1):
        properties = shape._inline.docPr
        if not properties.get("descr") or not properties.get("title"):
            raise RuntimeError(f"inline figure {index} lacks alt text or title")
    if len(document.tables) < 10:
        raise RuntimeError(f"unexpectedly few evidence/design tables: {len(document.tables)}")
    for table_index, table in enumerate(document.tables, start=1):
        properties = table._tbl.tblPr
        width = properties.find(qn("w:tblW"))
        indent = properties.find(qn("w:tblInd"))
        columns = table._tbl.tblGrid.findall(qn("w:gridCol"))
        if width is None or width.get(qn("w:w")) != str(base.USABLE_WIDTH_DXA):
            raise RuntimeError(f"table {table_index} lacks explicit 9360-DXA width")
        if indent is None or indent.get(qn("w:w")) != str(base.TABLE_INDENT_DXA):
            raise RuntimeError(f"table {table_index} lacks explicit 120-DXA indent")
        if sum(int(column.get(qn("w:w"))) for column in columns) != base.USABLE_WIDTH_DXA:
            raise RuntimeError(f"table {table_index} grid does not total 9360 DXA")

    header_text = " ".join(
        paragraph.text
        for current_section in document.sections
        for paragraph in current_section.header.paragraphs
    )
    if "V6.1" not in header_text or re.search(r"V(?:5|7|8)", header_text):
        raise RuntimeError(f"running header is stale or ambiguous: {header_text}")
    footer_xml = " ".join(
        current_section.footer._element.xml for current_section in document.sections
    )
    if "PAGE" not in footer_xml:
        raise RuntimeError("retained PAGE footer field was lost")
    opening = "\n".join(paragraph.text for paragraph in document.paragraphs[:12])
    if "V6.1 视觉全细节融合研究方案" not in opening:
        raise RuntimeError("V6.1 cover title was not written")
    if "V6.1" not in document.core_properties.title:
        raise RuntimeError("V6.1 core-properties title was not written")

    return {
        "sections": len(document.sections),
        "headings": len(headings),
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "inline_figures": len(document.inline_shapes),
        "page_size_in": [round(page_width, 3), round(page_height, 3)],
        "margins_in": [round(value, 3) for value in margins],
        "different_first_page": bool(section.different_first_page_header_footer),
        **package_foundation(reference, path),
    }


def build_document(args: argparse.Namespace) -> dict[str, Any]:
    repo = args.repo.resolve()
    reference = args.reference.resolve()
    content_path = args.content.resolve()
    benchmark_path = args.benchmark.resolve()
    figures_dir = args.figures_dir.resolve()
    output = args.output.resolve()
    validate_inputs(reference, content_path, benchmark_path, figures_dir, output)

    content = content_path.read_text(encoding="utf-8")
    benchmark = json.loads(benchmark_path.read_text(encoding="utf-8"))
    output.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(reference, output)
    document = Document(output)
    base.clear_main_body(document)
    replace_header_version(document)
    base.set_update_fields(document)

    properties = document.core_properties
    properties.title = "岩心 RGB–NIR–SWIR V6.1 视觉全细节融合研究方案"
    properties.subject = "GeoCoreFusion V6.1 visual-full-detail fusion"
    properties.author = "GeoCoreFusion Project"
    properties.last_modified_by = "GeoCoreFusion Project"
    properties.keywords = (
        "drill core; RGB-NIR-SWIR; subpixel registration; denoised RGB detail; "
        "screened Poisson; visual full detail; observation back projection"
    )
    properties.comments = (
        "Built from the retained V5 visual template; Word/PDF/PNG visual QA required."
    )
    properties.created = datetime(2026, 7, 20)
    properties.modified = datetime(2026, 7, 20)
    properties.revision = 61

    add_cover(document, args.build_date)
    add_overview(document)
    state = base.BuildState(
        document=document,
        figures_dir=figures_dir,
        benchmark=benchmark,
        edition="V6.1",
        figure_specs=FIGURE_SPECS,
    )
    base.render_markdown(state, content)
    if missing := sorted(set(FIGURE_SPECS) - state.inserted_figures):
        heading = document.add_paragraph("补充图件", style="Heading 1")
        heading.paragraph_format.page_break_before = True
        base.set_keep(heading, next_=True)
        for filename in missing:
            base.add_figure(state, filename)
    add_traceability_appendices(state)
    document.save(output)

    if sha256_file(reference).lower() != REFERENCE_SHA256:
        raise RuntimeError("retained V5 reference changed during V6.1 authoring")
    audit = audit_document(output, reference, content)
    audit.update(
        {
            "repo": str(repo),
            "reference": str(reference),
            "reference_sha256": REFERENCE_SHA256,
            "content": str(content_path),
            "benchmark": str(benchmark_path),
            "figures_dir": str(figures_dir),
            "output": str(output),
            "output_sha256": sha256_file(output),
            "render_qa": "NOT_RUN_BY_BUILDER",
        }
    )
    return audit


def parse_args() -> argparse.Namespace:
    repo = Path(__file__).resolve().parents[1]
    parser = argparse.ArgumentParser(
        description="Build the V6.1 DOCX from the retained V5 template without rendering."
    )
    parser.add_argument("--repo", type=Path, default=repo)
    parser.add_argument("--reference", type=Path, default=DEFAULT_REFERENCE)
    parser.add_argument(
        "--content",
        type=Path,
        default=repo
        / "artifacts"
        / "v61_research"
        / "evidence"
        / "v61_document_content.md",
    )
    parser.add_argument(
        "--benchmark",
        type=Path,
        default=repo
        / "artifacts"
        / "v61_research"
        / "experiments"
        / "benchmark_summary.json",
    )
    parser.add_argument(
        "--figures-dir",
        type=Path,
        default=repo / "artifacts" / "v61_research" / "figures",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=repo
        / "artifacts"
        / "v61_research"
        / "deliverables"
        / OUTPUT_NAME,
    )
    parser.add_argument("--build-date", default="2026 年 7 月 20 日")
    return parser.parse_args()


def main() -> int:
    audit = build_document(parse_args())
    print(json.dumps(audit, ensure_ascii=False, indent=2))
    print(
        "\nIMPORTANT: DOCX construction and structural audit are complete, but "
        "delivery remains blocked until Word/PDF/PNG rendering and 100% page inspection pass."
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
