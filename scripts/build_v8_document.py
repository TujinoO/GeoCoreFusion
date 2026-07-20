"""Build the V8 research-route DOCX with the retained V5 visual system.

The 1,700-line Markdown/OOXML rendering engine lives in
``build_v6_document.py``.  This module supplies only the V8 edition contract:
cover and running-header text, figure metadata, Wald evidence appendices, input
validation, and structural audits.  The retained V5 file is copied before the
main body is replaced, so its styles, theme, numbering foundation,
header/footer parts, and section/page system remain the design authority.

The builder intentionally does not claim visual completion.  Its output must
still be exported by Microsoft Word, rasterized, and inspected page by page.
No file or directory is deleted by this script.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import shutil
import zipfile
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as _Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

try:  # Direct execution: ``python scripts/build_v8_document.py``.
    import build_v6_document as base
except ModuleNotFoundError:  # Module execution/import from the repository root.
    from scripts import build_v6_document as base


REFERENCE_SHA256 = base.REFERENCE_SHA256
DEFAULT_REFERENCE = Path(
    r"E:\GRP Docs\2026.07 岩心近景影像融合技术"
    r"\高水平论文研究路线与实验方案_V5最终结果补充版.docx"
)
OUTPUT_NAME = "高水平论文研究路线与实验方案_V8条件融合与证据闭环版.docx"


FIGURE_SPECS: dict[str, dict[str, Any]] = {
    "01_v8_conditional_route.png": {
        "title": "V8 条件融合总路线：可辨识性门控、受约束恢复与 fail-closed 输出",
        "alt": (
            "V8 从候选无关同 MTF 带通可辨识性检验开始，按可辨识、弱可辨识和不可辨识"
            "三类进入受约束恢复、保守恢复或关闭 RGB 注入，并经传感器前向闭环验收的流程图。"
        ),
    },
    "02_identifiability_spectral_map.png": {
        "title": "候选无关带通可辨识性光谱图：两场景的判定差异与冻结门限",
        "alt": (
            "3DSSZ 与 ZKH3 在 367 个光谱波段上的带通关系强度、空间分块交叉验证证据和"
            "可辨识性分类对比图。"
        ),
        "note": (
            "该门控不读取任何候选融合结果或 Wald 伪真值；当前两场景判定仅是方法可行性证据，"
            "必须冻结后在独立岩心、独立传感器标定和 HR-NIR/SWIR 真值上复验。"
        ),
    },
    "03_wald_cross_scene_methods.png": {
        "title": "严格 Wald 伪高分辨率基准：跨场景方法性能与负迁移",
        "alt": (
            "Bicubic、MTF-GLP、HPM、Brovey、RGB ridge 与 rank-8 低秩光谱锥方法在"
            "3DSSZ 和 ZKH3 上的 RMSE、SAM、ERGAS 和 2201 nm 指标比较图。"
        ),
        "note": (
            "Wald 降质提供同一数据内部的伪高分辨率真值，不等同于独立仪器获得的真实"
            "HR-SWIR；图中方法参数不得使用伪真值调优。"
        ),
    },
    "04_wald_2201_visual_truth.png": {
        "title": "2201 nm Wald 伪真值对照：细节恢复、伪纹理与暗区误差",
        "alt": (
            "两场景 2201 nm 的 Wald 伪高分辨率真值、Bicubic、传统细节注入、低秩 ridge"
            "和 0.5 度光谱锥约束结果及局部放大对照图。"
        ),
        "page_break_before": True,
    },
    "05_spectral_cone_geometry.png": {
        "title": "0.5° 光谱锥约束：保留公共明暗分量并截断光谱正交扰动",
        "alt": (
            "以双三次基准光谱为锥轴，将候选增量分解为平行公共明暗分量和正交光谱形状分量，"
            "再把后者限制在固定 0.5 度半角内的几何示意图。"
        ),
    },
    "06_3dssz_pareto_failure.png": {
        "title": "3DSSZ 真实 ROI 的 Pareto 失败：提高细节幅值同时放大非相干高频",
        "alt": (
            "17 至 19 组真实 3DSSZ 消融中相干细节幅值 beta、非相干高频 R_perp、"
            "前向误差和暗区指标的 Pareto 散点与失败区域。"
        ),
        "note": (
            "该图是无独立 HR-SWIR 真值条件下的同数据诊断；它支持停止盲目增大 strength，"
            "但不能单独判定真实高频恢复质量。"
        ),
    },
    "07_registration_uncertainty_sigma_points.png": {
        "title": "配准不确定性五点传播：从局部协方差代理到融合稳健性区间",
        "alt": (
            "RGB-NIR-SWIR 配准链、局部二维位移协方差代理、中心与四个 sigma 点扰动，"
            "以及融合指标区间传播的技术流程图。"
        ),
        "note": (
            "当前协方差来自局部匹配曲面和残差统计，是待标定的不确定性代理；"
            "在独立控制点或位移靶验证前不得解释为真实覆盖概率。"
        ),
    },
    "08_scientific_vs_visualization_contract.png": {
        "title": "双产品契约：科学融合立方体与 RGB 纹理化可视化严格分流",
        "alt": (
            "同一输入经过科学融合立方体和 RGB 纹理化显示两个输出分支，分别标注"
            "允许的光谱定量用途、视觉用途、元数据和禁止事项。"
        ),
    },
    "09_validation_counterfactual_matrix.png": {
        "title": "验证与反事实矩阵：Wald、RGB shuffle、warp sweep 和 RGB-only edge",
        "alt": (
            "按真值等级、扰动类型、预期响应、失败信号和可支持论文主张组织的验证矩阵图。"
        ),
    },
    "10_truth_acquisition_roadmap.png": {
        "title": "独立真值获取与论文冻结路线：从 MTF/噪声标定到 HR-SWIR 验收",
        "alt": (
            "分辨率靶与暗区阶跃靶、微位移超采样、局部 HR-NIR/SWIR 拍摄、"
            "冻结外部验证和论文主张升级的阶段路线图。"
        ),
        "page_break_before": True,
    },
}


OVERVIEW_ITEMS = [
    "目标分层：RGB 等效观感、观测一致性与独立 HR-NIR/SWIR 真值不可混同",
    "失败审计：V5/V6/V7 为何在暗区细节、非相干高频和跨场景泛化上受限",
    "条件 UARF-Cycle：候选无关可辨识性门控与 fail-closed 科学输出",
    "可辨识分支：rank-8 系数 ridge、0.5° 固定光谱锥与原生波段前向细化",
    "配准证据链：局部协方差代理、五点传播和独立控制点标定边界",
    "真实 ROI 与 Pareto 审计：停止盲扫公共 gain、局部残差和 additive strength",
    "严格 Wald 基准：跨场景收益、负迁移、暗区和 2201 nm 伪真值证据",
    "双产品契约：科学融合立方体与 RGB 纹理化可视化明确分流",
    "反事实与独立真值：RGB shuffle、warp sweep、边缘审计和真值采集",
    "论文贡献、消融矩阵、验收/停止准则以及代码—证据追溯路线",
]


CODE_MAP_ROWS = [
    ["模块", "主要文件", "V8 职责", "证据/验收入口"],
    [
        "可辨识性门控",
        "src/geocorefusion/identifiability.py",
        "同 MTF 带通关系、空间分块 CV、三态 fail-closed 判定",
        "v8_identifiability.csv / .md；tests/test_identifiability.py",
    ],
    [
        "受约束融合",
        "src/geocorefusion/fusion.py；lowrank.py",
        "rank-8 系数桥、丰度/残差原型、最终 L1 与单纯形约束",
        "v8_detail_sweep.csv；Wald 基准；tests/test_fusion.py",
    ],
    [
        "光谱保护",
        "src/geocorefusion/spectral_guard.py",
        "固定 0.5° 光谱锥、公共明暗分量保留、正交扰动截断",
        "tests/test_spectral_guard.py；Wald cone 审计字段",
    ],
    [
        "配准不确定性",
        "src/geocorefusion/registration_uncertainty.py",
        "局部协方差代理与五点 sigma-point 传播",
        "tests/test_registration_uncertainty.py；warp sweep",
    ],
    [
        "质量评价",
        "src/geocorefusion/quality.py",
        "原生幅值 beta、A、R_perp、暗区、flat 区、Edge F1 与方向",
        "v7/v8 机器结果；tests/test_quality_detail.py",
    ],
    [
        "严格 Wald",
        "scripts/run_v8_wald_benchmark.py",
        "MTF 降质、伪 HR 真值隔离、跨场景九方法比较",
        "v8_wald_benchmark.json；v8_wald_band_metrics.csv",
    ],
    [
        "条件决策",
        "src/geocorefusion/pipeline.py；validation.py",
        "gate off/on、科学/显示产品分流、反复制检查",
        "tests/test_validation.py；独立真值验收清单",
    ],
]


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def replace_header_version(document: _Document) -> None:
    """Change only visible running-header text, retaining its part/layout."""

    replacement = (
        "GeoCoreFusion V8    岩心 RGB-NIR-SWIR 条件融合、证据闭环与论文实验方案"
    )
    for section in document.sections:
        header = section.header
        changed = False
        for paragraph in header.paragraphs:
            text_nodes = list(paragraph._p.iter(qn("w:t")))
            visible_text = "".join(node.text or "" for node in text_nodes)
            if re.search(r"V(?:5|6|7)", visible_text) and text_nodes:
                text_nodes[0].text = replacement
                for text_node in text_nodes[1:]:
                    text_node.text = ""
                changed = True
        if changed:
            continue

        paragraph = next((p for p in header.paragraphs if p.text.strip()), None)
        if paragraph is None:
            raise RuntimeError("retained running header has no editable text paragraph")
        if paragraph.runs:
            paragraph.runs[0].text = replacement
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(replacement)


def add_cover(document: _Document, build_date: str) -> None:
    spacer = document.add_paragraph(style="Normal")
    spacer.paragraph_format.space_after = Pt(54)

    eyebrow = document.add_paragraph(style="Normal")
    eyebrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eyebrow.paragraph_format.space_after = Pt(16)
    run = eyebrow.add_run("RESEARCH ROADMAP & EXPERIMENTAL PLAN")
    base.set_run_font(run, latin="Calibri", size=10.5, color=base.BLUE, bold=True)

    title = document.add_paragraph(style="Normal")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    title.paragraph_format.line_spacing = 1.04
    run = title.add_run("岩心 RGB–NIR–SWIR 近景影像配准与融合")
    base.set_run_font(run, latin="Calibri", size=28, color=base.NAVY, bold=True)
    run = title.add_run("\nV8 条件融合与证据闭环研究方案")
    base.set_run_font(run, latin="Calibri", size=24, color=base.NAVY, bold=True)

    subtitle = document.add_paragraph(style="Normal")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    run = subtitle.add_run(
        "候选无关可辨识性 · 受约束光谱恢复 · fail-closed 科学输出 · 独立真值路线"
    )
    base.set_run_font(run, latin="Calibri", size=14, color=base.MID_BLUE)

    project = document.add_paragraph(style="Normal")
    project.alignment = WD_ALIGN_PARAGRAPH.CENTER
    project.paragraph_format.space_after = Pt(7)
    run = project.add_run("GeoCoreFusion | RGB–NIR–SWIR Drill-Core Image Fusion")
    base.set_run_font(run, latin="Calibri", size=11, color=base.GRAY, bold=True)

    date_paragraph = document.add_paragraph(style="Normal")
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_paragraph.add_run(f"{build_date} | V8 条件融合与证据闭环版")
    base.set_run_font(run, latin="Calibri", size=10.5, color=base.GRAY)
    date_paragraph.paragraph_format.space_after = Pt(36)

    status = document.add_paragraph(style="Normal")
    status.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = status.add_run(
        "条件结论：3DSSZ 关闭科学 RGB 注入 · ZKH3 为两场景策略证据 · 独立真值待补齐"
    )
    base.set_run_font(run, latin="Calibri", size=9.5, color=base.ORANGE, bold=True)
    status.paragraph_format.space_after = Pt(8)

    document.add_paragraph(style="Normal").add_run().add_break(WD_BREAK.PAGE)


def add_overview(document: _Document) -> None:
    heading = document.add_paragraph("内容概览", style="Heading 1")
    base.set_keep(heading, next_=True)
    lead = document.add_paragraph(
        "V8 不再把“肉眼接近 RGB”当作 SWIR 高频真实恢复证据，而是先判断目标波段的"
        "RGB—NIR/SWIR 带通关系是否可辨识，再决定恢复、保守恢复或关闭注入。"
        "全文始终区分真实 ROI 诊断、严格 Wald 伪真值与独立仪器真值。",
        style="Normal",
    )
    lead.paragraph_format.space_after = Pt(7)

    numbering_id = base.new_numbering_instance(document)
    for text in OVERVIEW_ITEMS:
        paragraph = document.add_paragraph(style="List Number")
        base.apply_numbering(paragraph, numbering_id)
        paragraph.paragraph_format.space_after = Pt(3)
        base.add_rich_text(
            paragraph,
            text,
            base_size=10.5,
            base_color=base.MID_BLUE,
        )
        for run in paragraph.runs:
            run.bold = True

    toc_heading = document.add_paragraph(style="Normal")
    run = toc_heading.add_run("自动目录")
    base.set_run_font(run, latin="Calibri", size=13, color=base.BLUE, bold=True)
    toc_heading.paragraph_format.space_before = Pt(9)
    toc_heading.paragraph_format.space_after = Pt(4)
    base.set_keep(toc_heading, next_=True)

    toc = document.add_paragraph(style="Normal")
    toc.paragraph_format.space_after = Pt(4)
    base.add_complex_field(
        toc,
        r'TOC \o "1-2" \h \z \u',
        "目录域将在 Microsoft Word 打开或打印时更新。",
        font_size=9,
        color=base.GRAY,
    )
    note = document.add_paragraph(
        "提示：V8 允许因证据表和图件增加页数；最终页码以 Word 字段更新和逐页质检为准。",
        style="Figure Caption",
    )
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    document.add_paragraph(style="Normal").add_run().add_break(WD_BREAK.PAGE)


def _method_rows(benchmark: dict[str, Any]) -> list[list[str]]:
    rows = [
        [
            "场景",
            "方法",
            "全谱 RMSE",
            "SAM / °",
            "ERGAS",
            "LR forward RMSE",
            "RMSE 相对 Bicubic",
        ]
    ]
    order = benchmark.get("method_order", [])
    for scene in benchmark.get("scenes", []):
        methods = scene.get("methods", {})
        for method_name in order:
            method = methods.get(method_name)
            if not method:
                continue
            metrics = method["evaluation"]["full_spectrum"]
            rows.append(
                [
                    str(scene.get("scene", "")).upper(),
                    str(method.get("label", method_name)),
                    f"{metrics['rmse']:.6f}",
                    f"{metrics['sam_mean_deg']:.4f}",
                    f"{metrics['ergas']:.4f}",
                    f"{metrics['lr_forward_rmse']:.6f}",
                    f"{metrics['rmse_improvement_vs_bicubic_pct']:+.2f}%",
                ]
            )
    return rows


def _policy_rows(benchmark: dict[str, Any]) -> list[list[str]]:
    rows = [
        [
            "场景",
            "门控",
            "选择方法",
            "全谱 RMSE",
            "全谱改善",
            "2201 nm RMSE",
            "2201 nm 改善",
            "锥截断比例",
        ]
    ]
    decisions = benchmark.get("fail_closed_policy", {}).get("decisions", {})
    for scene_name, decision in decisions.items():
        values = decision.get("selected_evaluation", {})
        rows.append(
            [
                scene_name.upper(),
                str(decision.get("gate_state", "unknown")).upper(),
                str(decision.get("selected_method", "未记录")),
                f"{values.get('full_spectrum_rmse', float('nan')):.6f}",
                (
                    f"{values.get('full_spectrum_rmse_improvement_vs_bicubic_pct', float('nan')):+.2f}%"
                ),
                f"{values.get('2201nm_rmse', float('nan')):.6f}",
                f"{values.get('2201nm_rmse_improvement_vs_bicubic_pct', float('nan')):+.2f}%",
                f"{values.get('cone_clip_fraction', float('nan')):.3f}",
            ]
        )
    return rows


def _target_band_rows(benchmark: dict[str, Any]) -> list[list[str]]:
    rows = [
        [
            "场景",
            "方法",
            "波段",
            "RMSE",
            "PSNR / dB",
            "rho",
            "beta",
            "A",
            "R_perp",
        ]
    ]
    selected = {
        key: value.get("selected_method")
        for key, value in benchmark.get("fail_closed_policy", {})
        .get("decisions", {})
        .items()
    }
    for scene in benchmark.get("scenes", []):
        scene_name = str(scene.get("scene", ""))
        method_name = selected.get(scene_name)
        method = scene.get("methods", {}).get(method_name, {})
        bands = method.get("evaluation", {}).get("target_bands", {})
        for band_name in ("901nm", "1651nm", "2201nm"):
            band = bands.get(band_name)
            if not band:
                continue
            detail = band.get("detail", {})
            rows.append(
                [
                    scene_name.upper(),
                    str(method.get("label", method_name)),
                    band_name,
                    f"{band['rmse']:.6f}",
                    f"{band['psnr_db']:.3f}",
                    f"{detail.get('rho', float('nan')):.3f}",
                    f"{detail.get('beta', float('nan')):.3f}",
                    f"{detail.get('energy_ratio_A', float('nan')):.3f}",
                    f"{detail.get('artifact_R_perp', float('nan')):.3f}",
                ]
            )
    return rows


def _protocol_rows(benchmark: dict[str, Any]) -> list[list[str]]:
    protocol = benchmark.get("protocol", {})
    return [
        ["协议项", "冻结值", "证据含义"],
        ["空间降质比例", str(protocol.get("ratio", "未记录")), "伪 HR 到 LR 的固定比例"],
        [
            "Nyquist MTF",
            str(protocol.get("nyquist_mtf", "未记录")),
            "同一 Wald 退化核口径",
        ],
        [
            "blocked CV",
            f"{protocol.get('blocked_cv_folds', '未记录')} folds；buffer={protocol.get('blocked_cv_buffer_lr', '未记录')} LR px",
            "降低空间泄漏，不读取伪 HR-HSI 真值调参",
        ],
        [
            "光谱锥半角",
            f"{protocol.get('spectral_cone_half_angle_deg', '未记录')}°",
            "预注册保护参数；truth_tuned=false",
        ],
        [
            "估计可访问数据",
            ", ".join(protocol.get("estimation_access", [])),
            "训练/选择阶段允许读取",
        ],
        [
            "仅评价可访问数据",
            ", ".join(protocol.get("evaluation_only_access", [])),
            "只用于最终评分，禁止方法选择",
        ],
    ]


def add_traceability_appendices(state: base.BuildState) -> None:
    heading = state.document.add_paragraph(
        "附录 C：V8 代码—配置—证据地图",
        style="Heading 1",
    )
    heading.paragraph_format.page_break_before = True
    base.set_keep(heading, next_=True)
    paragraph = state.document.add_paragraph(
        "下表把条件融合的每项论文主张追溯到实现、机器结果和测试入口。"
        "路径均相对于 GeoCoreFusion 仓库根目录。",
        style="Normal",
    )
    paragraph.paragraph_format.space_after = Pt(6)
    base.add_table(state, CODE_MAP_ROWS, caption="V8 代码—配置—证据追溯地图")

    heading = state.document.add_paragraph(
        "附录 D：严格 Wald 与 fail-closed 机器快照",
        style="Heading 1",
    )
    heading.paragraph_format.page_break_before = True
    base.set_keep(heading, next_=True)
    base.add_callout(
        state.document,
        "来源",
        (
            f"本附录在构建时读取 {state.benchmark.get('benchmark', 'v8_wald_benchmark')} "
            f"schema v{state.benchmark.get('schema_version', '未记录')}。"
            "主文静态数字若与本附录冲突，以机器 JSON、复现实验日志及其哈希为准。"
        ),
    )
    base.add_table(state, _protocol_rows(state.benchmark), caption="严格 Wald 冻结协议")
    base.add_table(state, _method_rows(state.benchmark), caption="跨场景九方法全谱 Wald 结果")
    base.add_table(state, _policy_rows(state.benchmark), caption="当前两场景 fail-closed 决策快照")
    base.add_table(
        state,
        _target_band_rows(state.benchmark),
        caption="门控所选方法的目标波段 Wald 细节诊断",
    )
    claim_boundary = state.benchmark.get("fail_closed_policy", {}).get(
        "claim_boundary",
        "当前策略仍需独立数据复验。",
    )
    base.add_callout(
        state.document,
        "强制证据边界",
        (
            "Wald 是严格的同数据伪真值协议，但不是独立 HR-SWIR 仪器真值。"
            "真实 ROI 的 RGB 同源细节指标也不是 SWIR 高频真值。"
            f"策略记录：{claim_boundary}"
        ),
        warning=True,
    )


def validate_inputs(
    reference: Path,
    content: Path,
    benchmark: Path,
    figures_dir: Path,
    output: Path,
) -> None:
    for path in (reference, content, benchmark):
        if not path.is_file():
            raise FileNotFoundError(path)
    if not figures_dir.is_dir():
        raise NotADirectoryError(figures_dir)

    missing = [name for name in FIGURE_SPECS if not (figures_dir / name).is_file()]
    if missing:
        raise FileNotFoundError(f"missing V8 figures: {missing}")
    if reference.resolve() == output.resolve():
        raise ValueError("output must be different from retained V5 reference")

    actual_hash = sha256_file(reference).lower()
    if actual_hash != REFERENCE_SHA256:
        raise RuntimeError(
            "retained V5 SHA-256 mismatch; re-distill before authoring: "
            f"expected {REFERENCE_SHA256}, got {actual_hash}"
        )

    benchmark_data = json.loads(benchmark.read_text(encoding="utf-8"))
    if benchmark_data.get("benchmark") != "strict_internal_wald_pseudo_hr":
        raise ValueError("benchmark is not the strict V8 internal Wald artifact")
    if not benchmark_data.get("scenes") or not benchmark_data.get("method_order"):
        raise ValueError("V8 Wald benchmark is missing scenes or method_order")

    markdown = content.read_text(encoding="utf-8")
    referenced = {
        Path(match).name
        for match in re.findall(r"!\[[^\]]*\]\(([^)\s]+\.png)", markdown)
    }
    unknown = sorted(referenced - set(FIGURE_SPECS))
    if unknown:
        raise ValueError(f"content references figures without V8 metadata: {unknown}")


def _assert_package_foundation(reference: Path, output: Path) -> dict[str, int]:
    required_parts = {
        "word/styles.xml",
        "word/theme/theme1.xml",
        "word/fontTable.xml",
        "word/numbering.xml",
        "word/header1.xml",
        "word/footer1.xml",
        "word/settings.xml",
        "word/webSettings.xml",
    }
    with zipfile.ZipFile(reference) as source_zip, zipfile.ZipFile(output) as final_zip:
        source_names = set(source_zip.namelist())
        final_names = set(final_zip.namelist())
        missing_from_source = sorted(required_parts - source_names)
        missing_from_final = sorted(required_parts - final_names)
        if missing_from_source:
            raise RuntimeError(f"retained V5 lacks expected package parts: {missing_from_source}")
        if missing_from_final:
            raise RuntimeError(f"V8 lost retained package parts: {missing_from_final}")

        source_custom = {name for name in source_names if name.startswith("customXml/")}
        lost_custom = sorted(source_custom - final_names)
        if lost_custom:
            raise RuntimeError(f"V8 lost retained customXml parts: {lost_custom}")

        source_media = {name for name in source_names if name.startswith("word/media/")}
        retained_media = source_media & final_names
        if retained_media != source_media:
            lost_media = sorted(source_media - retained_media)
            raise RuntimeError(f"V8 lost retained media parts: {lost_media}")

    return {
        "required_foundation_parts": len(required_parts),
        "retained_custom_xml_parts": len(source_custom),
        "retained_source_media_parts": len(source_media),
    }


def _expected_markdown_headings(markdown: str) -> list[str]:
    headings: list[tuple[int, str]] = []
    for line in markdown.splitlines():
        match = re.match(r"^(#{1,3})\s+(.+?)\s*$", line.strip())
        if match:
            headings.append((len(match.group(1)), match.group(2).strip()))
    for index, (level, _) in enumerate(headings):
        if level == 1:
            headings.pop(index)
            break
    return [text for _, text in headings]


def audit_built_document(
    path: Path,
    *,
    reference: Path,
    content_markdown: str,
) -> dict[str, Any]:
    document = Document(path)
    if len(document.sections) != 1:
        raise RuntimeError(f"expected one retained section, found {len(document.sections)}")
    section = document.sections[0]
    if section.orientation != WD_ORIENT.PORTRAIT:
        raise RuntimeError("V8 document is not portrait")

    page_width = section.page_width / 914400
    page_height = section.page_height / 914400
    if abs(page_width - 8.5) > 0.01 or abs(page_height - 11.0) > 0.01:
        raise RuntimeError(f"unexpected page size: {page_width:.3f} x {page_height:.3f} in")
    margins = [
        section.top_margin / 914400,
        section.right_margin / 914400,
        section.bottom_margin / 914400,
        section.left_margin / 914400,
    ]
    if any(abs(value - 1.0) > 0.01 for value in margins):
        raise RuntimeError(f"retained one-inch margins changed: {margins}")
    if not section.different_first_page_header_footer:
        raise RuntimeError("different-first-page header/footer setting was lost")

    heading_paragraphs = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.style and paragraph.style.name.startswith("Heading ")
    ]
    heading_texts = [paragraph.text.strip() for paragraph in heading_paragraphs]
    expected_counts = Counter(_expected_markdown_headings(content_markdown))
    actual_counts = Counter(heading_texts)
    missing_headings = list((expected_counts - actual_counts).elements())
    if missing_headings:
        raise RuntimeError(f"Markdown headings were not emitted as real Word headings: {missing_headings}")
    if len(heading_paragraphs) < 12:
        raise RuntimeError(f"unexpectedly few real Word headings: {len(heading_paragraphs)}")

    if len(document.inline_shapes) != len(FIGURE_SPECS):
        raise RuntimeError(
            f"expected {len(FIGURE_SPECS)} inline V8 figures, "
            f"found {len(document.inline_shapes)}"
        )
    for index, shape in enumerate(document.inline_shapes, start=1):
        doc_properties = shape._inline.docPr
        if not doc_properties.get("descr") or not doc_properties.get("title"):
            raise RuntimeError(f"inline figure {index} lacks title or alt text")

    if len(document.tables) < 5:
        raise RuntimeError(f"unexpectedly few evidence/design tables: {len(document.tables)}")
    for table_index, table in enumerate(document.tables, start=1):
        table_properties = table._tbl.tblPr
        table_width = table_properties.find(qn("w:tblW"))
        table_indent = table_properties.find(qn("w:tblInd"))
        grid_columns = table._tbl.tblGrid.findall(qn("w:gridCol"))
        grid_widths = [int(column.get(qn("w:w"))) for column in grid_columns]
        if table_width is None or table_width.get(qn("w:w")) != str(base.USABLE_WIDTH_DXA):
            raise RuntimeError(f"table {table_index} lacks explicit 9360-DXA tblW")
        if table_indent is None or table_indent.get(qn("w:w")) != str(base.TABLE_INDENT_DXA):
            raise RuntimeError(f"table {table_index} lacks explicit 120-DXA tblInd")
        if sum(grid_widths) != base.USABLE_WIDTH_DXA:
            raise RuntimeError(f"table {table_index} grid does not total 9360 DXA")
        for row_index, row in enumerate(table.rows, start=1):
            if len(row.cells) != len(grid_widths):
                raise RuntimeError(
                    f"table {table_index} row {row_index} has merged/incomplete geometry"
                )
            for column_index, (cell, expected_width) in enumerate(
                zip(row.cells, grid_widths),
                start=1,
            ):
                cell_width = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
                if (
                    cell_width is None
                    or cell_width.get(qn("w:type")) != "dxa"
                    or int(cell_width.get(qn("w:w"))) != expected_width
                ):
                    raise RuntimeError(
                        f"table {table_index} row {row_index} column {column_index} "
                        "does not match the explicit DXA grid"
                    )

    header_text = " ".join(
        paragraph.text
        for current_section in document.sections
        for paragraph in current_section.header.paragraphs
    )
    if "V8" not in header_text or re.search(r"V(?:5|6|7)", header_text):
        raise RuntimeError(f"running header version is stale or ambiguous: {header_text}")
    footer_xml = " ".join(
        current_section.footer._element.xml for current_section in document.sections
    )
    if "PAGE" not in footer_xml:
        raise RuntimeError("retained PAGE footer field was lost")
    if "V8 条件融合与证据闭环研究方案" not in "\n".join(
        paragraph.text for paragraph in document.paragraphs[:12]
    ):
        raise RuntimeError("V8 cover title was not written")
    if "V8" not in document.core_properties.title:
        raise RuntimeError("V8 core-properties title was not written")

    style_names = {style.name for style in document.styles}
    required_styles = {
        "Normal",
        "Heading 1",
        "Heading 2",
        "Heading 3",
        "Figure Caption",
        "Header",
        "Footer",
    }
    if missing_styles := sorted(required_styles - style_names):
        raise RuntimeError(f"retained Word styles were lost: {missing_styles}")

    package_audit = _assert_package_foundation(reference, path)
    return {
        "sections": len(document.sections),
        "headings": len(heading_paragraphs),
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "inline_figures": len(document.inline_shapes),
        "page_size_in": [round(page_width, 3), round(page_height, 3)],
        "margins_in": [round(value, 3) for value in margins],
        "different_first_page": bool(section.different_first_page_header_footer),
        "page_count_policy": "UNBOUNDED_INCREASE_ALLOWED; verify after Word render",
        **package_audit,
    }


def build_document(args: argparse.Namespace) -> dict[str, Any]:
    repo = args.repo.resolve()
    reference = args.reference.resolve()
    content_path = args.content.resolve()
    benchmark_path = args.benchmark.resolve()
    figures_dir = args.figures_dir.resolve()
    output = args.output.resolve()

    validate_inputs(reference, content_path, benchmark_path, figures_dir, output)
    content = content_path.read_text(encoding="utf-8")
    benchmark = json.loads(benchmark_path.read_text(encoding="utf-8"))

    output.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(reference, output)
    document = Document(output)
    base.clear_main_body(document)
    replace_header_version(document)
    base.set_update_fields(document)

    properties = document.core_properties
    properties.title = "岩心 RGB–NIR–SWIR V8 条件融合与证据闭环研究方案"
    properties.subject = "GeoCoreFusion V8 conditional UARF-Cycle and evidence closure"
    properties.author = "GeoCoreFusion Project"
    properties.last_modified_by = "GeoCoreFusion Project"
    properties.keywords = (
        "drill core; RGB-NIR-SWIR; identifiability; fail-closed fusion; "
        "Wald protocol; spectral cone; registration uncertainty"
    )
    properties.comments = (
        "Built from the retained V5 visual template; Word/PDF/PNG visual QA required."
    )
    properties.created = datetime(2026, 7, 20)
    properties.modified = datetime(2026, 7, 20)
    properties.revision = 8

    add_cover(document, args.build_date)
    add_overview(document)
    state = base.BuildState(
        document=document,
        figures_dir=figures_dir,
        benchmark=benchmark,
        edition="V8",
        figure_specs=FIGURE_SPECS,
    )
    base.render_markdown(state, content)

    missing_after_render = sorted(set(FIGURE_SPECS) - state.inserted_figures)
    if missing_after_render:
        heading = document.add_paragraph("补充图件", style="Heading 1")
        heading.paragraph_format.page_break_before = True
        base.set_keep(heading, next_=True)
        for filename in missing_after_render:
            base.add_figure(state, filename)

    add_traceability_appendices(state)
    document.save(output)

    if sha256_file(reference).lower() != REFERENCE_SHA256:
        raise RuntimeError("retained V5 reference changed during V8 authoring")
    audit = audit_built_document(
        output,
        reference=reference,
        content_markdown=content,
    )
    audit.update(
        {
            "repo": str(repo),
            "reference": str(reference),
            "reference_sha256": REFERENCE_SHA256,
            "content": str(content_path),
            "benchmark": str(benchmark_path),
            "figures_dir": str(figures_dir),
            "output": str(output),
            "output_sha256": sha256_file(output),
            "render_qa": "NOT_RUN_BY_BUILDER",
        }
    )
    return audit


def parse_args() -> argparse.Namespace:
    script_repo = Path(__file__).resolve().parents[1]
    parser = argparse.ArgumentParser(
        description="Build the V8 DOCX from the retained V5 template without rendering."
    )
    parser.add_argument("--repo", type=Path, default=script_repo)
    parser.add_argument("--reference", type=Path, default=DEFAULT_REFERENCE)
    parser.add_argument(
        "--content",
        type=Path,
        default=script_repo
        / "artifacts"
        / "v8_research"
        / "evidence"
        / "v8_document_content.md",
    )
    parser.add_argument(
        "--benchmark",
        type=Path,
        default=script_repo
        / "artifacts"
        / "v7_research"
        / "evidence"
        / "v8_wald_benchmark.json",
    )
    parser.add_argument(
        "--figures-dir",
        type=Path,
        default=script_repo / "artifacts" / "v8_research" / "figures",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=script_repo
        / "artifacts"
        / "v8_research"
        / "deliverables"
        / OUTPUT_NAME,
    )
    parser.add_argument("--build-date", default="2026 年 7 月 20 日")
    return parser.parse_args()


def main() -> int:
    audit = build_document(parse_args())
    print(json.dumps(audit, ensure_ascii=False, indent=2))
    print(
        "\nIMPORTANT: DOCX construction and structural audit are complete, but "
        "delivery remains blocked until Microsoft Word -> PDF -> PNG rendering "
        "and 100% inspection of every page pass."
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
