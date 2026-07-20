"""Build the V6 research-route DOCX from the retained V5 visual template.

This builder deliberately does not render Word/PDF.  The caller must complete
the Word -> PDF -> PNG render-and-inspect loop before the document is delivered.

The retained V5 DOCX is treated as read-only design authority.  The script:

* verifies its SHA-256 before and after authoring;
* copies it to a distinct output path;
* clears only the copied main-document body while retaining the final sectPr;
* preserves the source styles, theme, numbering, headers, footers and page
  system;
* rebuilds the body from v6_document_content.md;
* reads benchmark_summary.json for the machine-derived result appendix;
* inserts all ten V6 figures as inline images with Word SEQ captions;
* applies explicit DXA geometry to every generated table.

No file or directory is deleted by this script.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import shutil
from copy import deepcopy
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable, Sequence

from PIL import Image
from docx import Document
from docx.document import Document as _Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run


REFERENCE_SHA256 = (
    "8a50037e89c3d36674be87eb38b74b2a0ffed1670d96f2abcf2dd200b1dc657b"
)
DEFAULT_REFERENCE = Path(
    r"E:\GRP Docs\2026.07 岩心近景影像融合技术"
    r"\高水平论文研究路线与实验方案_V5最终结果补充版.docx"
)
OUTPUT_NAME = "高水平论文研究路线与实验方案_V6配准与融合优化版.docx"
USABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120

NAVY = "17365D"
BLUE = "2E74B5"
MID_BLUE = "1F4D78"
LIGHT_BLUE = "D9EAF7"
PALE_BLUE = "EDF4F9"
ORANGE = "E67E22"
PALE_ORANGE = "FFF1E6"
GRAY = "5F6978"
LIGHT_GRAY = "F3F5F7"
GRID_GRAY = "B8C4CE"
TEXT = "20252B"
WHITE = "FFFFFF"


FIGURE_SPECS: dict[str, dict[str, Any]] = {
    "01_registration_v6_route.png": {
        "title": "V6 跨模态配准：从坐标约定到独立 TRE 验收的闭环路线",
        "alt": "RGB、NIR、SWIR 跨模态配准的坐标一致、有效域、控制点、低阶形变和独立 TRE 验收流程图。",
    },
    "02_fusion_v6_route.png": {
        "title": "V6 SAMI-HPM：暗部相对对比度恢复与传感器观测闭环",
        "alt": "SAMI-HPM 从光谱统一、材料子空间、log-RGB 暗部特征到指数增益和观测回投影的技术路线。",
    },
    "03_registration_synthetic_accuracy.png": {
        "title": "受控合成真值组件试验上的 TRE/EPE 与证据边界",
        "alt": "粗仿射、ROI 仿射和稠密残差的受控合成 TRE 与 EPE 结果图。",
        "note": (
            "证据单位为 analysis-grid pixels。该试验是算法组件的受控真值试验："
            "注册 ECC/搜索阈值与真实生产配置不同，因此不能外推为真实岩心或生产配置已达到亚像素。"
        ),
    },
    "04_fusion_real_metrics.png": {
        "title": "两组真实 ROI 的 V5-matched 与 V6 整包方法对照",
        "alt": "3DSSZ 和 ZKH3 的观测网格自一致性及暗部结构传递指标对照图。",
        "note": (
            "同一场景的 V5/V6 上游产物哈希一致，构成 matched 对照；但融合多项设置同时改变，"
            "只能称为整包方法对照，不能称为单因素消融。continuous-cube 指标是 model-based "
            "observation-grid self-consistency，不是最终 HR 输出路径的严格 forward residual，"
            "因为 clip 与退化算子不交换。"
        ),
    },
    "05_3dssz_v5_v6_visual_comparison.png": {
        "title": "3DSSZ：V5-matched 与 V6 的全图及暗部局部结构对比",
        "alt": "3DSSZ 真实 ROI 中 V5-matched 与 V6 的全图、暗部和选定波段局部对比。",
        "page_break_before": True,
        "note": (
            "每幅 PNG 独立采用 2–98% 拉伸，仅用于结构可见性检查，不能比较绝对对比度或辐射保真。"
            "暗区 log-HF 相关由 RGB 同源驱动并以 RGB 结构评价，只是结构传递诊断，不是独立 SWIR 高频真值。"
        ),
    },
    "06_zkh3_v5_v6_visual_comparison.png": {
        "title": "ZKH3：V5-matched 与 V6 的全图及暗部局部结构对比",
        "alt": "ZKH3 真实 ROI 中 V5-matched 与 V6 的全图、暗部和选定波段局部对比。",
        "page_break_before": True,
        "note": (
            "每幅 PNG 独立采用 2–98% 拉伸，仅用于结构可见性检查，不能比较绝对对比度或辐射保真。"
            "图示不支持“无损恢复 SWIR 高频”的结论。"
        ),
    },
    "07_classical_pansharpening_vs_v6.png": {
        "title": "经典强度替换的视觉优势与跨谱段真实性边界",
        "alt": "GS、PC、HSV、Brovey、NNDiffuse 与 V6 的目标、视觉效果及证据边界对照图。",
        "note": (
            "经典方法接近 PAN/RGB 的视觉锐度主要来自强度替换或调制；"
            "它可作为空间清晰度基线，但不是 SWIR 辐射真实性上限。"
        ),
    },
    "08_evidence_ladder.png": {
        "title": "从视觉对齐到真实仪器亚像素与融合真实性的证据阶梯",
        "alt": "配准与融合从视觉检查、同数据分数、合成真值、标准靶标到独立光谱和地质真值的证据层级。",
    },
    "09_paper_experiment_matrix.png": {
        "title": "从算法真值、真实仪器到地质应用的论文实验矩阵",
        "alt": "涵盖合成数据、标准靶标、公共数据、真实岩心、独立光谱真值和地质下游任务的实验矩阵。",
    },
    "10_research_roadmap.png": {
        "title": "V6 从算法修正、真值闭环到高水平论文的阶段路线",
        "alt": "V6 基线冻结、真实亚像素证据、融合消融、安全性、独立真值和论文成稿路线图。",
    },
}


CODE_MAP_ROWS = [
    [
        "模块",
        "主要文件",
        "V6 职责",
        "证据/验证入口",
    ],
    [
        "配准核心",
        "src/geocorefusion/registration.py",
        "端点坐标、有效掩膜、无效边界、边界峰拒绝、TRE/EPE",
        "tests/test_registration.py；合成 TRE/EPE",
    ],
    [
        "融合核心",
        "src/geocorefusion/fusion.py",
        "log-RGB、local-SNR、材料系数回归、指数增益、双重回投影",
        "tests/test_fusion.py；matched/合成融合",
    ],
    [
        "质量评价",
        "src/geocorefusion/quality.py",
        "901/1651/2201 nm 的 RGB 定义最暗 20% log-HF 结构诊断",
        "benchmark_summary.json；真实 ROI 指标",
    ],
    [
        "流程编排",
        "src/geocorefusion/pipeline.py",
        "配准、融合、质量输出和 NaN 安全预览链条",
        "runs/*_v6；src/geocorefusion/validation.py",
    ],
    [
        "配置模型",
        "src/geocorefusion/config.py",
        "confidence mode、gain back-projection 等 V6 参数",
        "配置解析测试与运行日志",
    ],
    [
        "V6 配置",
        "configs/*_roi_fusion_v6.yaml",
        "rank=12、variational、gain strength=0.20 的真实 ROI 配置",
        "3DSSZ/ZKH3 两场景",
    ],
    [
        "matched 基线",
        "configs/*_roi_fusion_v5_matched.yaml",
        "共享修正后配准几何的 V5-style 对照",
        "上游产物哈希一致性审计",
    ],
    [
        "可复现基准",
        "scripts/run_v6_benchmarks.py",
        "合成真值与真实 matched 指标汇总",
        "artifacts/v6_research/experiments/benchmark_summary.json",
    ],
]


OVERVIEW_ITEMS = [
    "证据口径：已证明、已支持与待验证的结论边界",
    "配准诊断：坐标、边界和评价为何会污染亚像素判断",
    "V6 配准：粗仿射—双向控制点—受限推扫形变—独立 TRE",
    "暗部机理：亮度门控与线性增益造成的双重细节压制",
    "SAMI-HPM：log 相对对比度、材料自适应与观测闭环",
    "真实 matched 结果：暗部结构收益与观测网格自一致性代价",
    "经典方法定位：视觉接近 RGB 不等于 SWIR 高频真实",
    "论文实验：合成、靶标、公共数据、真实岩心和独立光谱真值",
    "代码地图、消融、预注册门槛和可复现配置",
    "创新点、风险边界、文献证据与分阶段实施路线",
]


@dataclass
class BuildState:
    document: _Document
    figures_dir: Path
    benchmark: dict[str, Any]
    # V6 remains the default so the existing CLI keeps its behavior.  Later
    # editions can reuse this rendering engine without firing V6-only hooks.
    edition: str = "V6"
    figure_specs: dict[str, dict[str, Any]] | None = None
    figure_number: int = 0
    table_number: int = 0
    equation_number: int = 0
    inserted_figures: set[str] = field(default_factory=set)
    current_heading: str = ""
    numbered_list_id: int | None = None


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run_font(
    run: Run,
    *,
    latin: str | None = None,
    east_asia: str | None = None,
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    if latin:
        run.font.name = latin
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        rfonts.set(qn("w:ascii"), latin)
        rfonts.set(qn("w:hAnsi"), latin)
    if east_asia:
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        rfonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_keep(paragraph: Paragraph, *, next_: bool = False, together: bool = False) -> None:
    paragraph.paragraph_format.keep_with_next = next_
    paragraph.paragraph_format.keep_together = together
    paragraph.paragraph_format.widow_control = True


def shade_paragraph(paragraph: Paragraph, fill: str, border: str) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)

    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    left = pbdr.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        pbdr.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), border)


def add_callout(
    document: _Document,
    label: str,
    text: str,
    *,
    warning: bool = False,
) -> Paragraph:
    p = document.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.08
    shade_paragraph(p, PALE_ORANGE if warning else PALE_BLUE, ORANGE if warning else BLUE)
    label_run = p.add_run(f"{label}｜")
    set_run_font(
        label_run,
        latin="Calibri",
        size=9.5,
        color=ORANGE if warning else MID_BLUE,
        bold=True,
    )
    text_run = p.add_run(text)
    set_run_font(text_run, latin="Calibri", size=9.5, color=TEXT)
    set_keep(p, together=True)
    return p


def clear_main_body(document: _Document) -> None:
    body = document._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def replace_header_version(document: _Document) -> None:
    for section in document.sections:
        for paragraph in section.header.paragraphs:
            changed = False
            for text_node in paragraph._p.iter(qn("w:t")):
                if text_node.text and "V5" in text_node.text:
                    text_node.text = text_node.text.replace("V5", "V6")
                    changed = True
            if changed:
                continue
            if paragraph.text.strip():
                for run in paragraph.runs:
                    run.text = ""
                paragraph.add_run(
                    "GeoCoreFusion V6    岩心 RGB-NIR-SWIR 近景影像配准与融合研究路线及实验方案"
                )


def set_update_fields(document: _Document) -> None:
    settings = document.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def add_complex_field(
    paragraph: Paragraph,
    instruction: str,
    cached_text: str,
    *,
    font_size: float | None = None,
    color: str | None = None,
) -> None:
    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    begin_run._r.append(begin)

    instr_run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    instr_run._r.append(instr)

    separate_run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)

    result = paragraph.add_run(cached_text)
    set_run_font(result, latin="Calibri", size=font_size, color=color)

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def add_cover(document: _Document, build_date: str) -> None:
    spacer = document.add_paragraph(style="Normal")
    spacer.paragraph_format.space_after = Pt(54)

    eyebrow = document.add_paragraph(style="Normal")
    eyebrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eyebrow.paragraph_format.space_after = Pt(16)
    r = eyebrow.add_run("RESEARCH ROADMAP & EXPERIMENTAL PLAN")
    set_run_font(r, latin="Calibri", size=10.5, color=BLUE, bold=True)

    title = document.add_paragraph(style="Normal")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    title.paragraph_format.line_spacing = 1.04
    r = title.add_run("岩心 RGB–NIR–SWIR 近景影像配准与融合")
    set_run_font(r, latin="Calibri", size=28, color=NAVY, bold=True)
    r = title.add_run("\nV6 研究路线与高水平论文实验方案")
    set_run_font(r, latin="Calibri", size=24, color=NAVY, bold=True)

    subtitle = document.add_paragraph(style="Normal")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    r = subtitle.add_run("坐标一致亚像素证据 · 阴影稳健 SAMI-HPM · 传感器观测闭环")
    set_run_font(r, latin="Calibri", size=14, color=MID_BLUE)

    project = document.add_paragraph(style="Normal")
    project.alignment = WD_ALIGN_PARAGRAPH.CENTER
    project.paragraph_format.space_after = Pt(7)
    r = project.add_run("GeoCoreFusion | RGB–NIR–SWIR Drill-Core Image Fusion")
    set_run_font(r, latin="Calibri", size=11, color=GRAY, bold=True)

    date_p = document.add_paragraph(style="Normal")
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = date_p.add_run(f"{build_date} | V6 配准与融合优化版")
    set_run_font(r, latin="Calibri", size=10.5, color=GRAY)
    date_p.paragraph_format.space_after = Pt(36)

    status = document.add_paragraph(style="Normal")
    status.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = status.add_run("证据分级：合成真值已证明 · 真实 matched 已支持 · 独立真值待补齐")
    set_run_font(r, latin="Calibri", size=9.5, color=ORANGE, bold=True)
    status.paragraph_format.space_after = Pt(8)

    p = document.add_paragraph(style="Normal")
    p.add_run().add_break(WD_BREAK.PAGE)


def add_overview(document: _Document) -> None:
    h = document.add_paragraph("内容概览", style="Heading 1")
    set_keep(h, next_=True)
    lead = document.add_paragraph(
        "本版将“视觉效果”改写为可证伪的几何精度、暗部结构、观测一致性和独立真值问题。"
        "建议先阅读证据口径与摘要，再按配准、融合、实验、风险和实施路线深入。",
        style="Normal",
    )
    lead.paragraph_format.space_after = Pt(7)

    for index, text in enumerate(OVERVIEW_ITEMS, start=1):
        p = document.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Inches(0.12)
        p.paragraph_format.space_after = Pt(3)
        n = p.add_run(f"{index:02d}  ")
        set_run_font(n, latin="Calibri", size=10.5, color=ORANGE, bold=True)
        r = p.add_run(text)
        set_run_font(r, latin="Calibri", size=10.5, color=MID_BLUE, bold=True)

    toc_h = document.add_paragraph(style="Normal")
    toc_run = toc_h.add_run("自动目录")
    set_run_font(toc_run, latin="Calibri", size=13, color=BLUE, bold=True)
    toc_h.paragraph_format.space_before = Pt(9)
    toc_h.paragraph_format.space_after = Pt(4)
    set_keep(toc_h, next_=True)
    toc = document.add_paragraph(style="Normal")
    toc.paragraph_format.space_after = Pt(4)
    add_complex_field(
        toc,
        r'TOC \o "1-2" \h \z \u',
        "目录域将在 Microsoft Word 打开或打印时更新。",
        font_size=9,
        color=GRAY,
    )
    note = document.add_paragraph(
        "提示：最终页码以完成 Word 字段更新并通过逐页 PNG 质检后的版本为准。",
        style="Figure Caption",
    )
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p = document.add_paragraph(style="Normal")
    p.add_run().add_break(WD_BREAK.PAGE)


def remove_markdown_marks(text: str) -> str:
    return (
        text.replace(r"\(", "")
        .replace(r"\)", "")
        .replace("**", "")
        .replace("__", "")
        .replace(chr(96), "")
    )


def add_rich_text(
    paragraph: Paragraph,
    text: str,
    *,
    base_size: float | None = None,
    base_color: str | None = None,
) -> None:
    code_tick = re.escape(chr(96))
    pattern = re.compile(
        rf"(\*\*.+?\*\*|{code_tick}.+?{code_tick}|\$[^$]+\$|\\\(.+?\\\))"
    )
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            r = paragraph.add_run(text[cursor : match.start()])
            set_run_font(r, latin="Calibri", size=base_size, color=base_color)
        token = match.group(0)
        if token.startswith("**"):
            r = paragraph.add_run(token[2:-2])
            set_run_font(
                r,
                latin="Calibri",
                size=base_size,
                color=base_color,
                bold=True,
            )
        elif token.startswith(chr(96)):
            r = paragraph.add_run(token[1:-1])
            set_run_font(
                r,
                latin="Consolas",
                east_asia="微软雅黑",
                size=(base_size or 10.5) - 0.5,
                color=MID_BLUE,
            )
        else:
            math_text = token[1:-1] if token.startswith("$") else token[2:-2]
            r = paragraph.add_run(latex_to_linear(math_text))
            set_run_font(
                r,
                latin="Cambria Math",
                east_asia="Cambria Math",
                size=base_size,
                color=base_color,
            )
        cursor = match.end()
    if cursor < len(text):
        r = paragraph.add_run(text[cursor:])
        set_run_font(r, latin="Calibri", size=base_size, color=base_color)


def _extract_group(text: str, brace_index: int) -> tuple[str, int] | None:
    if brace_index >= len(text) or text[brace_index] != "{":
        return None
    level = 0
    for index in range(brace_index, len(text)):
        if text[index] == "{":
            level += 1
        elif text[index] == "}":
            level -= 1
            if level == 0:
                return text[brace_index + 1 : index], index + 1
    return None


def _replace_command_one_arg(text: str, command: str, formatter) -> str:
    needle = "\\" + command
    while True:
        start = text.find(needle)
        if start < 0:
            return text
        brace = start + len(needle)
        while brace < len(text) and text[brace].isspace():
            brace += 1
        group = _extract_group(text, brace)
        if group is None:
            text = text[:start] + command + text[start + len(needle) :]
            continue
        value, end = group
        text = text[:start] + formatter(latex_to_linear(value)) + text[end:]


def _replace_frac(text: str) -> str:
    needle = r"\frac"
    while True:
        start = text.find(needle)
        if start < 0:
            return text
        first_start = start + len(needle)
        while first_start < len(text) and text[first_start].isspace():
            first_start += 1
        first = _extract_group(text, first_start)
        if first is None:
            return text.replace(needle, "/", 1)
        numerator, first_end = first
        second_start = first_end
        while second_start < len(text) and text[second_start].isspace():
            second_start += 1
        second = _extract_group(text, second_start)
        if second is None:
            return text.replace(needle, "/", 1)
        denominator, second_end = second
        replacement = f"({latex_to_linear(numerator)})/({latex_to_linear(denominator)})"
        text = text[:start] + replacement + text[second_end:]


def latex_to_linear(text: str) -> str:
    value = text.strip().replace("\n", " ")
    value = value.replace(r"\begin{bmatrix}", "[").replace(r"\end{bmatrix}", "]")
    value = value.replace(r"\begin{matrix}", "[").replace(r"\end{matrix}", "]")
    value = value.replace(r"\left", "").replace(r"\right", "")
    value = re.sub(r"\\mathcal\s+([A-Za-z])", r"\1", value)
    value = re.sub(r"\\bar\s+([A-Za-z])", lambda m: f"{m.group(1)}̄", value)
    value = re.sub(r"\\hat\s+([A-Za-z])", lambda m: f"{m.group(1)}̂", value)
    value = re.sub(r"\\tilde\s+([A-Za-z])", lambda m: f"{m.group(1)}̃", value)
    value = _replace_frac(value)
    value = _replace_command_one_arg(value, "sqrt", lambda x: f"√({x})")
    for command in (
        "operatorname",
        "mathrm",
        "text",
        "mathbf",
        "mathit",
        "mathcal",
    ):
        value = _replace_command_one_arg(value, command, lambda x: x)
    value = _replace_command_one_arg(value, "bar", lambda x: f"{x}̄")
    value = _replace_command_one_arg(value, "hat", lambda x: f"{x}̂")
    value = _replace_command_one_arg(value, "tilde", lambda x: f"{x}̃")

    replacements = {
        r"\alpha": "α",
        r"\beta": "β",
        r"\gamma": "γ",
        r"\Delta": "Δ",
        r"\delta": "δ",
        r"\eta": "η",
        r"\varepsilon": "ε",
        r"\epsilon": "ε",
        r"\Lambda": "Λ",
        r"\lambda": "λ",
        r"\mu": "μ",
        r"\Pi": "Π",
        r"\rho": "ρ",
        r"\Sigma": "Σ",
        r"\sigma": "σ",
        r"\tau": "τ",
        r"\theta": "θ",
        r"\Omega": "Ω",
        r"\operatorname": "",
        r"\arg": "arg ",
        r"\exp": "exp",
        r"\log": "log",
        r"\max": "max",
        r"\min": "min",
        r"\sum": "Σ",
        r"\cap": "∩",
        r"\in": "∈",
        r"\approx": "≈",
        r"\leq": "≤",
        r"\le": "≤",
        r"\geq": "≥",
        r"\ge": "≥",
        r"\rightarrow": "→",
        r"\to": "→",
        r"\mid": "│",
        r"\ast": "∗",
        r"\cdot": "·",
        r"\times": "×",
        r"\odot": "⊙",
        r"\top": "ᵀ",
        r"\qquad": "  ",
        r"\quad": " ",
        r"\;": " ",
        r"\,": " ",
        r"\!": "",
        r"\{": "{",
        r"\}": "}",
        r"\|": "‖",
        r"\\": " ; ",
        "&": " ",
    }
    for old, new in sorted(replacements.items(), key=lambda item: len(item[0]), reverse=True):
        value = value.replace(old, new)
    remaining_commands = sorted(set(re.findall(r"\\[A-Za-z]+", value)))
    if remaining_commands:
        raise ValueError(
            "unhandled LaTeX command(s): " + ", ".join(remaining_commands)
        )
    if "\\" in value:
        raise ValueError(f"unhandled LaTeX escape in equation: {value}")
    value = re.sub(r"\s+", " ", value)
    value = value.replace("{", "(").replace("}", ")")
    return value.strip()


def add_equation(document: _Document, latex: str, number: int) -> Paragraph:
    p = document.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.right_indent = Inches(0.22)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.05
    text = latex_to_linear(latex)
    r = p.add_run(text)
    set_run_font(
        r,
        latin="Cambria Math",
        east_asia="Cambria Math",
        size=10,
        color=TEXT,
    )
    n = p.add_run(f"    ({number})")
    set_run_font(n, latin="Cambria Math", size=9, color=GRAY)
    set_keep(p, together=True)
    return p


def set_repeat_table_header(row) -> None:
    trpr = row._tr.get_or_add_trPr()
    header = trpr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        trpr.append(header)
    header.set(qn("w:val"), "true")


def set_table_row_cant_split(row) -> None:
    """Keep a logical table row on one page when Word paginates it."""
    trpr = row._tr.get_or_add_trPr()
    cant_split = trpr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        trpr.append(cant_split)
    cant_split.set(qn("w:val"), "true")


def set_cell_shading(cell: _Cell, fill: str) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_cell_margins(
    cell: _Cell,
    *,
    top: int = 90,
    start: int = 120,
    bottom: int = 90,
    end: int = 120,
) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    tcmar = tcpr.first_child_found_in("w:tcMar")
    if tcmar is None:
        tcmar = OxmlElement("w:tcMar")
        tcpr.append(tcmar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcmar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tcmar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table: Table) -> None:
    tblpr = table._tbl.tblPr
    borders = tblpr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblpr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "5" if edge.startswith("inside") else "7")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), GRID_GRAY)


def set_table_geometry(table: Table, widths: Sequence[int]) -> None:
    if len(widths) != len(table.columns):
        raise ValueError("width count must match table column count")
    if sum(widths) != USABLE_WIDTH_DXA:
        raise ValueError("table widths must total exactly 9360 DXA")

    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblpr = table._tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.insert(0, tblw)
    tblw.set(qn("w:w"), str(USABLE_WIDTH_DXA))
    tblw.set(qn("w:type"), "dxa")

    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tblind.set(qn("w:type"), "dxa")

    layout = tblpr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblpr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(width))
            tcw.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def content_widths(rows: Sequence[Sequence[str]]) -> list[int]:
    column_count = len(rows[0])
    scores: list[float] = []
    for col in range(column_count):
        lengths: list[float] = []
        for row in rows[:24]:
            text = remove_markdown_marks(row[col]) if col < len(row) else ""
            score = sum(1.55 if ord(ch) > 127 else 0.9 for ch in text)
            lengths.append(min(max(score, 4.0), 34.0))
        scores.append(max(7.0, min(30.0, sum(lengths) / max(len(lengths), 1))))

    minimum = 720 if column_count >= 5 else 900
    widths = [max(minimum, round(USABLE_WIDTH_DXA * score / sum(scores))) for score in scores]
    total = sum(widths)
    if total != USABLE_WIDTH_DXA:
        adjustable = sorted(range(column_count), key=lambda i: widths[i], reverse=True)
        delta = USABLE_WIDTH_DXA - total
        index = 0
        while delta:
            col = adjustable[index % len(adjustable)]
            step = 1 if delta > 0 else -1
            if widths[col] + step >= minimum:
                widths[col] += step
                delta -= step
            index += 1
    return widths


def is_short_or_numeric(text: str) -> bool:
    clean = remove_markdown_marks(text).strip()
    if len(clean) <= 12:
        return True
    return bool(
        re.fullmatch(
            r"[+\-−]?[0-9.,/%°≤≥–—\s]+(?:nm|px|倍|项|个|°)?",
            clean,
        )
    )


def add_caption(
    state: BuildState,
    label: str,
    number: int,
    title: str,
    *,
    keep_next: bool = False,
) -> Paragraph:
    p = state.document.add_paragraph(style="Figure Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(f"{label} ")
    set_run_font(r, latin="Calibri", size=9, color=GRAY, bold=True)
    add_complex_field(
        p,
        f"SEQ {'Figure' if label == '图' else 'Table'} \\* ARABIC",
        str(number),
        font_size=9,
        color=GRAY,
    )
    r = p.add_run(f"  {title}")
    set_run_font(r, latin="Calibri", size=9, color=GRAY)
    set_keep(p, next_=keep_next, together=True)
    return p


def table_caption_title(
    heading: str,
    rows: Sequence[Sequence[str]],
    *,
    edition: str = "V6",
) -> str:
    clean_heading = re.sub(r"^\d+(?:\.\d+)*\s*", "", heading).strip()
    header = " ".join(rows[0])
    if "波段" in header and "log-HF" in header:
        return f"{clean_heading}：暗部 log-HF 结构传递诊断"
    if "V5-matched" in header and "指标" in header:
        return f"{clean_heading}：观测网格自一致性"
    if "阶段" in header and "P95" in header:
        return "确定性合成真值配准误差"
    return clean_heading or f"{edition} 结果与实验设计"


def add_table(
    state: BuildState,
    rows: Sequence[Sequence[str]],
    *,
    caption: str,
) -> Table:
    if not rows or not rows[0]:
        raise ValueError("cannot create an empty table")
    col_count = len(rows[0])
    normalized = [
        list(row[:col_count]) + [""] * max(0, col_count - len(row))
        for row in rows
    ]

    state.table_number += 1
    add_caption(state, "表", state.table_number, caption, keep_next=True)

    table = state.document.add_table(rows=len(normalized), cols=col_count)
    try:
        table.style = "Normal Table"
    except KeyError:
        table.style = "Table Grid"
    set_table_geometry(table, content_widths(normalized))
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])

    body_size = 8.2 if col_count >= 5 else 8.7
    for row_index, (row, values) in enumerate(zip(table.rows, normalized)):
        set_table_row_cant_split(row)
        for col_index, (cell, value) in enumerate(zip(row.cells, values)):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                set_cell_shading(cell, MID_BLUE)
            elif row_index % 2 == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            else:
                set_cell_shading(cell, WHITE)

            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            if row_index == 0 or is_short_or_numeric(value):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_rich_text(
                p,
                value,
                base_size=8.4 if row_index == 0 else body_size,
                base_color=WHITE if row_index == 0 else TEXT,
            )
            if row_index == 0:
                for run in p.runs:
                    run.bold = True
            set_keep(p, together=False)

    after = state.document.add_paragraph(style="Normal")
    after.paragraph_format.space_after = Pt(2)
    return table


def add_figure(state: BuildState, filename: str) -> None:
    if filename in state.inserted_figures:
        return
    figure_specs = state.figure_specs if state.figure_specs is not None else FIGURE_SPECS
    spec = figure_specs.get(filename)
    if spec is None:
        raise KeyError(f"missing figure specification for {filename}")
    path = state.figures_dir / filename
    if not path.is_file():
        raise FileNotFoundError(path)

    if spec.get("page_break_before"):
        pbreak = state.document.add_paragraph(style="Normal")
        pbreak.add_run().add_break(WD_BREAK.PAGE)

    with Image.open(path) as image:
        width_px, height_px = image.size
    aspect = width_px / height_px
    max_width_in = float(spec.get("max_width_in", 6.20))
    max_height_in = float(spec.get("max_height_in", 6.65))
    width_in = min(max_width_in, max_height_in * aspect)

    p = state.document.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    inline = run.add_picture(str(path), width=Inches(width_in))
    inline._inline.docPr.set("descr", spec["alt"])
    inline._inline.docPr.set("title", spec["title"])
    set_keep(p, next_=True, together=True)

    state.figure_number += 1
    note = spec.get("note")
    add_caption(
        state,
        "图",
        state.figure_number,
        spec["title"],
        keep_next=bool(note),
    )
    if note:
        note_p = state.document.add_paragraph(style="Figure Caption")
        note_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        note_p.paragraph_format.left_indent = Inches(0.15)
        note_p.paragraph_format.right_indent = Inches(0.15)
        note_p.paragraph_format.space_after = Pt(7)
        r = note_p.add_run("图注边界｜")
        set_run_font(r, latin="Calibri", size=8.5, color=ORANGE, bold=True)
        r = note_p.add_run(note)
        set_run_font(r, latin="Calibri", size=8.5, color=GRAY)
        set_keep(note_p, together=True)

    state.inserted_figures.add(filename)


def base_numbering_id(document: _Document, style_name: str = "List Number") -> int:
    style = document.styles[style_name]
    numpr = style.element.pPr.numPr if style.element.pPr is not None else None
    if numpr is not None and numpr.numId is not None:
        return int(numpr.numId.val)
    return 5


def new_numbering_instance(document: _Document) -> int:
    numbering = document.part.numbering_part.element
    base_id = base_numbering_id(document)
    base_num = next(
        (
            num
            for num in numbering.findall(qn("w:num"))
            if int(num.get(qn("w:numId"))) == base_id
        ),
        None,
    )
    if base_num is None:
        base_num = next(iter(numbering.findall(qn("w:num"))), None)
    if base_num is None:
        return base_id
    abstract = base_num.find(qn("w:abstractNumId"))
    if abstract is None:
        return base_id
    base_abstract_id = int(abstract.get(qn("w:val")))
    base_abstract = next(
        (
            item
            for item in numbering.findall(qn("w:abstractNum"))
            if int(item.get(qn("w:abstractNumId"))) == base_abstract_id
        ),
        None,
    )
    if base_abstract is None:
        return base_id

    abstract_ids = [
        int(item.get(qn("w:abstractNumId")))
        for item in numbering.findall(qn("w:abstractNum"))
        if item.get(qn("w:abstractNumId"))
    ]
    new_abstract_id = max(abstract_ids, default=0) + 1
    new_abstract = deepcopy(base_abstract)
    new_abstract.set(qn("w:abstractNumId"), str(new_abstract_id))
    nsid = new_abstract.find(qn("w:nsid"))
    if nsid is not None:
        nsid.set(qn("w:val"), f"{new_abstract_id:08X}"[-8:])
    template = new_abstract.find(qn("w:tmpl"))
    if template is not None:
        template.set(qn("w:val"), f"{(new_abstract_id + 0x10000000):08X}"[-8:])
    first_level = next(
        (
            level
            for level in new_abstract.findall(qn("w:lvl"))
            if level.get(qn("w:ilvl")) == "0"
        ),
        None,
    )
    if first_level is not None:
        start = first_level.find(qn("w:start"))
        if start is None:
            start = OxmlElement("w:start")
            first_level.insert(0, start)
        start.set(qn("w:val"), "1")

    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(new_abstract)
    else:
        numbering.insert(numbering.index(first_num), new_abstract)

    existing = [
        int(num.get(qn("w:numId")))
        for num in numbering.findall(qn("w:num"))
        if num.get(qn("w:numId"))
    ]
    new_id = max(existing, default=0) + 1
    new_num = OxmlElement("w:num")
    new_num.set(qn("w:numId"), str(new_id))
    new_abs_ref = OxmlElement("w:abstractNumId")
    new_abs_ref.set(qn("w:val"), str(new_abstract_id))
    new_num.append(new_abs_ref)
    numbering.append(new_num)
    return new_id


def apply_numbering(paragraph: Paragraph, num_id: int, level: int = 0) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    numpr = ppr.find(qn("w:numPr"))
    if numpr is None:
        numpr = OxmlElement("w:numPr")
        ppr.append(numpr)
    ilvl = numpr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        numpr.append(ilvl)
    ilvl.set(qn("w:val"), str(level))
    numid = numpr.find(qn("w:numId"))
    if numid is None:
        numid = OxmlElement("w:numId")
        numpr.append(numid)
    numid.set(qn("w:val"), str(num_id))


def normalize_scientific_wording(markdown: str) -> str:
    replacement = (
        "691–2518 nm 的 nominal 5-nm 网格（367 个波段，精确包含 2518 nm 端点；"
        "尾部为 2511/2516/2518 nm，故并非全轴严格等间隔 5 nm）"
    )
    markdown = markdown.replace(
        "691–2518 nm、5 nm 间隔、367 个波段",
        replacement,
    )
    markdown = markdown.replace(
        "691-2518 nm、5 nm 间隔、367 个波段",
        replacement,
    )
    return markdown


def flush_paragraph(state: BuildState, lines: list[str]) -> None:
    if not lines:
        return
    text = " ".join(line.strip() for line in lines).strip()
    if not text:
        lines.clear()
        return
    p = state.document.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.12
    add_rich_text(p, text, base_size=10.5, base_color=TEXT)
    lines.clear()


def parse_table_rows(lines: Sequence[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        row = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        rows.append(row)
        index += 1
    if len(rows) < 2:
        return rows, index
    separator = rows[1]
    if not all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in separator):
        return rows, index
    return [rows[0], *rows[2:]], index


def heading_style_for(level: int, text: str) -> str:
    if text.startswith("附录"):
        return "Heading 1"
    if level == 1:
        return "Heading 1"
    if level == 2:
        return "Heading 2"
    return "Heading 3"


def add_heading_block(state: BuildState, level: int, text: str, first_body: bool) -> None:
    style = heading_style_for(level, text)
    p = state.document.add_paragraph(text, style=style)
    set_keep(p, next_=True)
    state.current_heading = text

    if state.edition != "V6":
        return

    if "摘要式核心结论" in text:
        reg = state.benchmark["registration_synthetic"]
        fusion = state.benchmark["fusion_synthetic"]
        add_callout(
            state.document,
            "机器结果快照",
            (
                f"合成粗仿射 TRE 中位数 {reg['coarse_affine_tre']['median_px']:.4f} px，"
                f"稠密残差 EPE P95 {reg['dense_residual_epe']['p95_px']:.4f} px；"
                f"合成暗纹理独立真值相关由 "
                f"{fusion['v5_style']['dark_log_detail_correlation']:.5f} 提升至 "
                f"{fusion['v6_intrinsic']['dark_log_detail_correlation']:.5f}。"
                "前者与后者均是受控算法组件试验，不等价于生产配置或真实岩心真值。"
            ),
        )
    elif "4.1 matched 对照原则" in text:
        add_callout(
            state.document,
            "可比性审计",
            (
                "每个场景的 V5-matched 与 V6 上游产物哈希一致，使用相同修正后配准几何；"
                "但 log 特征、置信模式、指数增益、回投影等多项同时改变，"
                "所以当前结果是整包方法对照，不是单因素消融。"
            ),
            warning=True,
        )
        add_callout(
            state.document,
            "指标口径",
            (
                "continuous_cube_observation 只表示 model-based observation-grid "
                "self-consistency；由于 clip 与退化不交换，它不是最终 HR 输出路径的严格 "
                "forward residual。暗区 log-HF 相关由 RGB 同源驱动和评价，只用于诊断结构传递。"
            ),
            warning=True,
        )
    elif "4.4 合成融合真值结果" in text:
        add_callout(
            state.document,
            "受控配置边界",
            (
                "合成融合使用 rank=3、bicubic 和 strength=0.28；真实 V6 使用 rank=12、"
                "variational 和 strength=0.20。合成暗纹理独立真值增益约为 "
                "r=0.47717→0.49877（+0.02160），这是当前最可防守的融合真值证据，"
                "但不是生产 V6 的端到端复现。"
            ),
            warning=True,
        )


def render_markdown(state: BuildState, markdown: str) -> None:
    markdown = normalize_scientific_wording(markdown)
    lines = markdown.splitlines()
    paragraph_lines: list[str] = []
    index = 0
    first_title_skipped = False
    first_body_heading = True
    in_math = False
    math_lines: list[str] = []

    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()

        if in_math:
            if stripped.endswith("$$"):
                closing = stripped[:-2].strip()
                if closing:
                    math_lines.append(closing)
                state.equation_number += 1
                add_equation(
                    state.document,
                    "\n".join(math_lines),
                    state.equation_number,
                )
                math_lines = []
                in_math = False
            else:
                math_lines.append(raw)
            index += 1
            continue

        if stripped.startswith("$$"):
            flush_paragraph(state, paragraph_lines)
            remainder = stripped[2:]
            if remainder.endswith("$$") and len(remainder) > 2:
                state.equation_number += 1
                add_equation(
                    state.document,
                    remainder[:-2],
                    state.equation_number,
                )
            else:
                in_math = True
                if remainder:
                    math_lines.append(remainder)
            state.numbered_list_id = None
            index += 1
            continue

        if not stripped:
            flush_paragraph(state, paragraph_lines)
            state.numbered_list_id = None
            index += 1
            continue

        if stripped == "---":
            flush_paragraph(state, paragraph_lines)
            index += 1
            continue

        heading_match = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading_match:
            flush_paragraph(state, paragraph_lines)
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            if not first_title_skipped and level == 1:
                first_title_skipped = True
            else:
                if text in {"文档定位与证据口径", "摘要式核心结论"}:
                    level = 1
                add_heading_block(state, level, text, first_body_heading)
                first_body_heading = False
            state.numbered_list_id = None
            index += 1
            continue

        if stripped.startswith("[建议插图："):
            flush_paragraph(state, paragraph_lines)
            filenames = re.findall(r"[0-9]{2}_[A-Za-z0-9_]+\.png", stripped)
            if not filenames:
                raise ValueError(f"could not parse figure placeholder: {stripped}")
            for filename in filenames:
                add_figure(state, filename)
            state.numbered_list_id = None
            index += 1
            continue

        markdown_image = re.fullmatch(
            r"!\[[^\]]*\]\(([^)\s]+\.png)(?:\s+[\"'][^\"']*[\"'])?\)",
            stripped,
        )
        if markdown_image:
            flush_paragraph(state, paragraph_lines)
            add_figure(state, Path(markdown_image.group(1)).name)
            state.numbered_list_id = None
            index += 1
            continue

        if stripped.startswith("|"):
            flush_paragraph(state, paragraph_lines)
            rows, next_index = parse_table_rows(lines, index)
            if len(rows) >= 2:
                add_table(
                    state,
                    rows,
                    caption=table_caption_title(
                        state.current_heading,
                        rows,
                        edition=state.edition,
                    ),
                )
                index = next_index
                state.numbered_list_id = None
                continue

        bullet_match = re.match(r"^[-*]\s+(.+)$", stripped)
        if bullet_match:
            flush_paragraph(state, paragraph_lines)
            # The retained template's first List Bullet style is linked to a
            # decimal numbering definition.  Its second bullet style is the
            # actual glyph-based bullet and preserves the source typography.
            p = state.document.add_paragraph(style="List Bullet 2")
            p.paragraph_format.space_after = Pt(2)
            add_rich_text(p, bullet_match.group(1), base_size=10.3, base_color=TEXT)
            set_keep(p)
            state.numbered_list_id = None
            index += 1
            continue

        number_match = re.match(r"^\d+[.)]\s+(.+)$", stripped)
        if number_match:
            flush_paragraph(state, paragraph_lines)
            if state.numbered_list_id is None:
                state.numbered_list_id = new_numbering_instance(state.document)
            p = state.document.add_paragraph(style="List Number")
            apply_numbering(p, state.numbered_list_id)
            p.paragraph_format.space_after = Pt(2)
            add_rich_text(p, number_match.group(1), base_size=10.3, base_color=TEXT)
            set_keep(p)
            index += 1
            continue

        if stripped.startswith(">"):
            flush_paragraph(state, paragraph_lines)
            p = state.document.add_paragraph(style="Quote")
            add_rich_text(p, stripped.lstrip("> ").strip(), base_size=10, base_color=GRAY)
            state.numbered_list_id = None
            index += 1
            continue

        paragraph_lines.append(stripped)
        index += 1

    flush_paragraph(state, paragraph_lines)
    if in_math and math_lines:
        raise ValueError("unterminated display-math block in markdown")


def build_registration_rows(benchmark: dict[str, Any]) -> list[list[str]]:
    reg = benchmark["registration_synthetic"]
    mapping = [
        ("粗仿射 TRE", reg["coarse_affine_tre"]),
        ("ROI 仿射 TRE", reg["roi_affine_tre"]),
        ("稠密残差 EPE", reg["dense_residual_epe"]),
    ]
    rows = [["阶段", "均值 / px", "中位数 / px", "P95 / px", "最大值 / px", "证据范围"]]
    for name, values in mapping:
        rows.append(
            [
                name,
                f"{values['mean_px']:.4f}",
                f"{values['median_px']:.4f}",
                f"{values['p95_px']:.4f}",
                f"{values['max_px']:.4f}",
                "确定性合成 analysis-grid 真值",
            ]
        )
    return rows


def build_synthetic_fusion_rows(benchmark: dict[str, Any]) -> list[list[str]]:
    fusion = benchmark["fusion_synthetic"]
    v5 = fusion["v5_style"]
    v6 = fusion["v6_intrinsic"]
    return [
        ["指标", "V5-style", "V6 intrinsic", "变化", "解释"],
        [
            "暗部 log 细节相关",
            f"{v5['dark_log_detail_correlation']:.6f}",
            f"{v6['dark_log_detail_correlation']:.6f}",
            f"{v6['dark_log_detail_correlation'] - v5['dark_log_detail_correlation']:+.6f}",
            "独立合成 HR 系数场上的结构真值证据",
        ],
        [
            "系数观测 RMSE",
            f"{v5['coefficient_observation_rmse']:.6f}",
            f"{v6['coefficient_observation_rmse']:.6f}",
            f"{v6['coefficient_observation_rmse'] - v5['coefficient_observation_rmse']:+.6f}",
            "越小越好",
        ],
        [
            "增益观测 RMSE",
            f"{v5['gain_observation_rmse']:.6f}",
            f"{v6['gain_observation_rmse']:.6f}",
            f"{v6['gain_observation_rmse'] - v5['gain_observation_rmse']:+.6f}",
            "越小越好",
        ],
    ]


def build_real_summary_rows(benchmark: dict[str, Any]) -> list[list[str]]:
    real = benchmark["real_roi_matched"]
    rows = [
        [
            "场景",
            "方法",
            "continuous cube RMSE",
            "SAM / °",
            "mean band CC",
            "coeff. RMSE",
            "gain LR RMSE",
        ]
    ]
    for scene in ("3dssz", "zkh3"):
        for method in ("v5_matched", "v6"):
            item = real[scene][method]
            rows.append(
                [
                    scene.upper(),
                    "V5-matched" if method == "v5_matched" else "V6",
                    f"{item['continuous_cube_rmse']:.6f}",
                    f"{item['sam_mean_deg']:.4f}",
                    f"{item['band_cc_mean']:.6f}",
                    f"{item['coefficient_rmse']:.6f}",
                    f"{item['gain_lowres_rmse']:.6f}",
                ]
            )
    return rows


def build_dark_detail_rows(benchmark: dict[str, Any]) -> list[list[str]]:
    real = benchmark["real_roi_matched"]
    rows = [["场景", "波段", "V5-matched", "V6", "绝对提升", "证据类型"]]
    for scene in ("3dssz", "zkh3"):
        v5 = real[scene]["v5_matched"]["dark_log_detail_correlation"]
        v6 = real[scene]["v6"]["dark_log_detail_correlation"]
        for wavelength in ("901.0nm", "1651.0nm", "2201.0nm"):
            rows.append(
                [
                    scene.upper(),
                    wavelength.replace(".0", " "),
                    f"{v5[wavelength]:.3f}",
                    f"{v6[wavelength]:.3f}",
                    f"{v6[wavelength] - v5[wavelength]:+.3f}",
                    "RGB 同源结构传递诊断",
                ]
            )
    return rows


def add_traceability_appendices(state: BuildState) -> None:
    h = state.document.add_paragraph("附录 C：代码—配置—证据地图", style="Heading 1")
    h.paragraph_format.page_break_before = True
    set_keep(h, next_=True)
    p = state.document.add_paragraph(
        "该地图用于把论文主张追溯到实现、配置、测试和机器可读结果。"
        "路径均相对于 GeoCoreFusion 仓库根目录。",
        style="Normal",
    )
    p.paragraph_format.space_after = Pt(6)
    add_table(state, CODE_MAP_ROWS, caption="V6 代码—配置—证据追溯地图")

    h = state.document.add_paragraph("附录 D：机器可读基准快照", style="Heading 1")
    h.paragraph_format.page_break_before = True
    set_keep(h, next_=True)
    add_callout(
        state.document,
        "来源",
        (
            f"以下表格由 benchmark_summary.json 在构建时生成；时间戳为 "
            f"{state.benchmark.get('generated_at', '未记录')}。"
            "主文静态数字与此处不一致时，以机器可读结果和复现实验日志为准。"
        ),
    )
    add_table(
        state,
        build_registration_rows(state.benchmark),
        caption="确定性合成配准基准快照",
    )
    add_table(
        state,
        build_synthetic_fusion_rows(state.benchmark),
        caption="确定性合成融合基准快照",
    )
    add_table(
        state,
        build_real_summary_rows(state.benchmark),
        caption="真实 ROI matched 整包方法观测网格自一致性",
    )
    add_table(
        state,
        build_dark_detail_rows(state.benchmark),
        caption="真实 ROI 最暗 20% 区域 log-HF 结构传递诊断",
    )
    add_callout(
        state.document,
        "强制证据边界",
        (
            "真实 ROI 指标使用修正后几何与 model-based observation-grid self-consistency，"
            "没有独立 HR-HSI 真值；暗区 HF 相关由 RGB 同源驱动和评价。"
            "因此本附录不证明生产 V6 在真实场景达到亚像素，也不证明无损恢复 SWIR 高频。"
        ),
        warning=True,
    )


def validate_inputs(
    reference: Path,
    content: Path,
    benchmark: Path,
    figures_dir: Path,
    output: Path,
) -> None:
    for path in (reference, content, benchmark):
        if not path.is_file():
            raise FileNotFoundError(path)
    if not figures_dir.is_dir():
        raise NotADirectoryError(figures_dir)
    missing = [name for name in FIGURE_SPECS if not (figures_dir / name).is_file()]
    if missing:
        raise FileNotFoundError(f"missing V6 figures: {missing}")
    if reference.resolve() == output.resolve():
        raise ValueError("output must be different from retained reference")
    actual = sha256_file(reference)
    if actual.lower() != REFERENCE_SHA256:
        raise RuntimeError(
            "retained V5 SHA-256 mismatch; re-distill before authoring: "
            f"expected {REFERENCE_SHA256}, got {actual}"
        )


def audit_built_document(path: Path) -> dict[str, Any]:
    document = Document(path)
    if len(document.sections) != 1:
        raise RuntimeError(f"expected one section, found {len(document.sections)}")
    section = document.sections[0]
    if section.orientation != WD_ORIENT.PORTRAIT:
        raise RuntimeError("final document is not portrait")
    page_width = section.page_width / 914400
    page_height = section.page_height / 914400
    if abs(page_width - 8.5) > 0.01 or abs(page_height - 11.0) > 0.01:
        raise RuntimeError(f"unexpected page size: {page_width:.3f} x {page_height:.3f} in")
    if not section.different_first_page_header_footer:
        raise RuntimeError("different-first-page header/footer setting was lost")

    headings = [
        p for p in document.paragraphs if p.style and p.style.name.startswith("Heading ")
    ]
    if len(headings) < 30:
        raise RuntimeError(f"unexpectedly few real Word headings: {len(headings)}")
    if len(document.inline_shapes) != len(FIGURE_SPECS):
        raise RuntimeError(
            f"expected {len(FIGURE_SPECS)} inline figures, found {len(document.inline_shapes)}"
        )
    if len(document.tables) < 12:
        raise RuntimeError(f"unexpectedly few result/design tables: {len(document.tables)}")
    header_text = " ".join(
        p.text for s in document.sections for p in s.header.paragraphs
    )
    if "V6" not in header_text or "V5" in header_text:
        raise RuntimeError(f"running header version was not updated cleanly: {header_text}")

    for index, table in enumerate(document.tables, start=1):
        tblpr = table._tbl.tblPr
        tblw = tblpr.find(qn("w:tblW"))
        tblind = tblpr.find(qn("w:tblInd"))
        grid = table._tbl.tblGrid.findall(qn("w:gridCol"))
        if tblw is None or tblw.get(qn("w:w")) != str(USABLE_WIDTH_DXA):
            raise RuntimeError(f"table {index} is missing 9360-DXA tblW")
        if tblind is None or tblind.get(qn("w:w")) != str(TABLE_INDENT_DXA):
            raise RuntimeError(f"table {index} is missing 120-DXA tblInd")
        if sum(int(col.get(qn("w:w"))) for col in grid) != USABLE_WIDTH_DXA:
            raise RuntimeError(f"table {index} grid width does not total 9360 DXA")

    return {
        "sections": len(document.sections),
        "headings": len(headings),
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "inline_figures": len(document.inline_shapes),
        "page_size_in": [round(page_width, 3), round(page_height, 3)],
        "different_first_page": bool(section.different_first_page_header_footer),
    }


def build_document(args: argparse.Namespace) -> dict[str, Any]:
    repo = args.repo.resolve()
    reference = args.reference.resolve()
    content_path = args.content.resolve()
    benchmark_path = args.benchmark.resolve()
    figures_dir = args.figures_dir.resolve()
    output = args.output.resolve()

    validate_inputs(
        reference,
        content_path,
        benchmark_path,
        figures_dir,
        output,
    )
    content = content_path.read_text(encoding="utf-8")
    benchmark = json.loads(benchmark_path.read_text(encoding="utf-8"))

    output.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(reference, output)
    document = Document(output)
    clear_main_body(document)
    replace_header_version(document)
    set_update_fields(document)

    properties = document.core_properties
    properties.title = "岩心 RGB–NIR–SWIR 近景影像配准与融合 V6：研究路线与论文实验方案"
    properties.subject = "GeoCoreFusion V6 registration and SAMI-HPM fusion"
    properties.author = "GeoCoreFusion Project"
    properties.last_modified_by = "GeoCoreFusion Project"
    properties.keywords = (
        "drill core; RGB-NIR-SWIR; multimodal registration; subpixel TRE; "
        "hyperspectral fusion; SAMI-HPM; observation consistency"
    )
    properties.comments = (
        "Built from the retained V5 visual template; requires Word/PDF/PNG visual QA."
    )
    properties.created = datetime(2026, 7, 20)
    properties.modified = datetime(2026, 7, 20)
    properties.revision = 6

    add_cover(document, args.build_date)
    add_overview(document)
    state = BuildState(document=document, figures_dir=figures_dir, benchmark=benchmark)
    render_markdown(state, content)

    missing_after_render = sorted(set(FIGURE_SPECS) - state.inserted_figures)
    if missing_after_render:
        h = document.add_paragraph("补充图件", style="Heading 1")
        h.paragraph_format.page_break_before = True
        set_keep(h, next_=True)
        for filename in missing_after_render:
            add_figure(state, filename)

    add_traceability_appendices(state)
    document.save(output)

    if sha256_file(reference).lower() != REFERENCE_SHA256:
        raise RuntimeError("retained reference changed during authoring")
    audit = audit_built_document(output)
    audit.update(
        {
            "repo": str(repo),
            "reference": str(reference),
            "reference_sha256": REFERENCE_SHA256,
            "content": str(content_path),
            "benchmark": str(benchmark_path),
            "figures_dir": str(figures_dir),
            "output": str(output),
            "output_sha256": sha256_file(output),
            "render_qa": "NOT_RUN_BY_BUILDER",
        }
    )
    return audit


def parse_args() -> argparse.Namespace:
    script_repo = Path(__file__).resolve().parents[1]
    parser = argparse.ArgumentParser(
        description="Build V6 DOCX from retained V5 template without rendering."
    )
    parser.add_argument("--repo", type=Path, default=script_repo)
    parser.add_argument("--reference", type=Path, default=DEFAULT_REFERENCE)
    parser.add_argument(
        "--content",
        type=Path,
        default=script_repo
        / "artifacts"
        / "v6_research"
        / "evidence"
        / "v6_document_content.md",
    )
    parser.add_argument(
        "--benchmark",
        type=Path,
        default=script_repo
        / "artifacts"
        / "v6_research"
        / "experiments"
        / "benchmark_summary.json",
    )
    parser.add_argument(
        "--figures-dir",
        type=Path,
        default=script_repo / "artifacts" / "v6_research" / "figures",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=script_repo
        / "artifacts"
        / "v6_research"
        / "deliverables"
        / OUTPUT_NAME,
    )
    parser.add_argument("--build-date", default="2026 年 7 月 20 日")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    audit = build_document(args)
    print(json.dumps(audit, ensure_ascii=False, indent=2))
    print(
        "\nIMPORTANT: DOCX construction is complete, but delivery is blocked until "
        "Microsoft Word -> PDF -> PNG rendering and 100% inspection of every page pass."
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
